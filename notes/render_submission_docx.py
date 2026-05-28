from __future__ import annotations

import re
import shutil
import textwrap
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image
from matplotlib import font_manager


ROOT = Path("/root/dev/pinn-reliability-paper")
NOTES = ROOT / "notes"
FIG_DIR = ROOT / "minimal_pinn" / "results" / "paper_figures" / "v1"
DOCX_PATH = Path("/root/dev/我们何时能信任数据_初稿_07_4.7.docx")
BACKUP_PATH = Path("/root/dev/我们何时能信任数据_初稿_07_4.7.before_full_final_sync_with_figures.docx")

MAIN_MD = NOTES / "paper_submission_polished_zh.md"
RESULTS_MD = NOTES / "results_section_zh.md"
STRATEGY_CSV = ROOT / "minimal_pinn" / "results" / "region_aware" / "region_aware_compare_v3_u3" / "strategy_summary.csv"
FIG6_PATH = FIG_DIR / "figure_06_region_aware_compare.png"
FIG5_PATH = FIG_DIR / "figure_05_recalibrated_dimension_combined.png"
FIG7_PATH = FIG_DIR / "figure_07_operator_to_topology.png"
FIG1_PATH = FIG_DIR / "figure_01_rel_l2_phase_maps.png"
FIG2_PATH = FIG_DIR / "figure_02_reliability_phase_maps.png"
FIG4_PATH = FIG_DIR / "figure_04_burgers_probability_composite.png"
FIG8_PATH = FIG_DIR / "figure_08_calibration_threshold_robustness.png"
FIG9_PATH = FIG_DIR / "figure_09_seed_statistical_uncertainty.png"


def configure_chinese_fonts() -> None:
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.otf",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            font_manager.fontManager.addfont(str(path))
            name = font_manager.FontProperties(fname=str(path)).get_name()
            plt.rcParams["font.sans-serif"] = [name, "DejaVu Sans"]
            plt.rcParams["axes.unicode_minus"] = False
            return


configure_chinese_fonts()
mpl.rcParams.update({
    "figure.facecolor": "white",
    "axes.facecolor": "white",
    "axes.spines.right": False,
    "axes.spines.top": False,
    "axes.linewidth": 0.8,
    "font.size": 8.5,
    "savefig.facecolor": "white",
})


FIGURES = [
    {
        "section": "### 6. 跨方程的可靠性相空间",
        "path": FIG1_PATH,
        "caption": "图 1. 四类方程在观测退化空间中的相对 L2 误差相图。泊松方程在当前扫描范围内未形成实用失效边界；定常 Stokes-Poiseuille 流的退化集中在低观测和中高噪声角落；Fisher-KPP 方程呈现较规则的前沿型过渡；黏性 Burgers 方程形成更宽、更不规则的高风险区域。",
    },
    {
        "section": "### 6. 跨方程的可靠性相空间",
        "path": FIG2_PATH,
        "caption": "图 2. 四类方程的 R_lin 标量化投影相图。R_lin 仅用于案例内排序和边界可视化；对失效来源的解释仍回到物理约束、训练轨迹、参考解误差和结构特征四维后验状态。",
    },
    {
        "section": "### 8. 系统依赖的边界语义",
        "path": FIG_DIR / "figure_03_regime_maps.png",
        "caption": "图 4. 规则边界与复杂边界的语义对比。该图基于细化矩阵与迁移实验中的经验分布关系绘制，用于概括定常 Stokes-Poiseuille 流与黏性 Burgers 方程在语义层上的差异。",
    },
    {
        "section": "### 9. 边界不确定性的统计证据",
        "path": FIG4_PATH,
        "caption": "图 5. 黏性 Burgers 方程临界带附近的多随机种子边界。左侧颜色表示 5 个随机种子中的越界率；中间给出 Wilson 区间宽度；右侧汇总不同噪声下的平均越界率。结果表明该边界更适合被视为具有统计宽度的过渡带。",
    },
    {
        "section": "### 9. 边界不确定性的统计证据",
        "path": FIG9_PATH,
        "caption": "图 6. 高密度多随机种子边界点的统计不确定性。每个点给出越界率及 95% Wilson 区间；稳定端点使用 30 个随机种子，临界点追加到 40 个随机种子。该图用于区分确定端点和宽过渡带，不解释为部署场景中的精确失效概率。",
    },
    {
        "section": "### 7. 为什么单一误差不足",
        "path": FIG5_PATH,
        "caption": "图 3. 单一误差指标无法解释失效来源。左侧显示各案例在四个诊断维度上的平均退化强度，中间给出相对 L2 误差对各维度的解释度，右侧统计每个案例的主导退化维度。黏性 Burgers 方程同时牵动物理约束、参考解误差和结构特征，因此更需要回到四维诊断状态。",
    },
    {
        "section": "### 10. 可迁移性与校准",
        "path": FIG8_PATH,
        "caption": "图 7. 校准拆分与阈值敏感性。左侧显示两类留出评估中的 rel_l2 与 R_lin 排序关系；右侧显示 1.25、1.5 与 2.0 倍基线阈值下各方程越界比例。阈值会移动边界位置，但不会消除方程间的主要语义差异。",
    },
    {
        "section": "### 11. 探索性外推：失效机制引导的训练干预",
        "path": FIG6_PATH,
        "caption": "图 8. 探索性训练干预具有系统依赖性。图中以基线为零点，比较不同区域感知策略对相对 L2 误差和 R_lin 的改变。该结果说明，诊断流程适合定位失效来源，但不能直接等同于通用训练改进策略。",
    },
    {
        "section": "### 12. 讨论",
        "path": FIG7_PATH,
        "caption": "图 9. 从方程算子性质到物理信息神经网络可靠性边界形态的经验解释。该图用于组织实验现象，不构成因果证明。",
    },
]


CASE_LABELS = {
    "poisson": "泊松方程",
    "stokes_poiseuille": "定常 Stokes-Poiseuille 流",
    "fisher_kpp": "Fisher-KPP 方程",
    "burgers": "黏性 Burgers 方程",
}

CASE_COLORS = {
    "poisson": "#2F7D5C",
    "stokes_poiseuille": "#3A6EA5",
    "fisher_kpp": "#C18B21",
    "burgers": "#B44B4B",
}

DIMENSION_LABELS = {
    "physics_consistency": "物理约束",
    "training_stability": "训练轨迹",
    "numerical_accuracy": "参考解误差",
    "structural_stability": "结构特征",
}


def save_png(fig: plt.Figure, path: Path, dpi: int = 320) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=dpi, bbox_inches="tight", pad_inches=0.06)
    plt.close(fig)


def load_matrix_data() -> pd.DataFrame:
    frames = [
        pd.read_csv(ROOT / "minimal_pinn" / "results" / "matrices" / "coarse_v1" / "matrix_summary.csv"),
        pd.read_csv(ROOT / "minimal_pinn" / "results" / "matrices" / "coarse_fisher_kpp_v1" / "matrix_summary.csv"),
    ]
    df = pd.concat(frames, ignore_index=True)
    df = df[df["case"].isin(CASE_LABELS)].copy()
    return df


def draw_matrix_heatmap(ax, df: pd.DataFrame, case: str, value: str, cmap: str, vmin=None, vmax=None, title: str = ""):
    sub = df[df["case"] == case].copy()
    piv = sub.pivot_table(index="num_observation", columns="noise_std", values=value, aggfunc="mean")
    piv = piv.sort_index(ascending=False).sort_index(axis=1)
    arr = piv.to_numpy(dtype=float)
    im = ax.imshow(arr, aspect="auto", cmap=cmap, vmin=vmin, vmax=vmax)
    ax.set_title(title or CASE_LABELS[case], fontsize=9.8, pad=5, fontweight="bold")
    ax.set_xticks(range(len(piv.columns)))
    ax.set_xticklabels([f"{x:.2f}" for x in piv.columns], rotation=45, ha="right", fontsize=7)
    ax.set_yticks(range(len(piv.index)))
    ax.set_yticklabels([str(int(x)) for x in piv.index], fontsize=7)
    ax.tick_params(length=0)
    for spine in ax.spines.values():
        spine.set_visible(True)
        spine.set_linewidth(0.5)
        spine.set_edgecolor("#D0D0D0")
    return im


def add_panel_label(ax, label: str) -> None:
    ax.text(-0.08, 1.04, label, transform=ax.transAxes, fontsize=10.5, fontweight="bold", va="bottom", ha="left")


def build_figure1() -> None:
    df = load_matrix_data()
    cases = ["poisson", "stokes_poiseuille", "fisher_kpp", "burgers"]
    vmin = max(df["rel_l2"].quantile(0.02), 1e-4)
    vmax = df["rel_l2"].quantile(0.98)

    fig, axes = plt.subplots(2, 2, figsize=(8.8, 7.2), constrained_layout=True)
    axes = axes.ravel()
    for ax, case in zip(axes, cases):
        im = draw_matrix_heatmap(ax, df, case, "rel_l2", "magma_r", vmin=vmin, vmax=vmax, title=CASE_LABELS[case])
        ax.set_xlabel("噪声标准差", fontsize=8)
        if ax in [axes[0], axes[2]]:
            ax.set_ylabel("观测点数", fontsize=8)
        else:
            ax.set_ylabel("")
        med = df[df["case"] == case]["rel_l2"].median()
        ax.text(0.03, 0.05, f"中位误差 {med:.3g}", transform=ax.transAxes, ha="left", va="bottom", fontsize=7.5,
                bbox=dict(facecolor="white", edgecolor="none", alpha=0.78, pad=1.5))
    for label, ax in zip("abcd", axes):
        add_panel_label(ax, label)
    cbar = fig.colorbar(im, ax=axes, location="right", shrink=0.82, pad=0.012)
    cbar.set_label("相对 L2 误差", fontsize=8)
    save_png(fig, FIG1_PATH)


def build_figure2() -> None:
    df = load_matrix_data()
    cases = ["poisson", "stokes_poiseuille", "fisher_kpp", "burgers"]
    fig, axes = plt.subplots(2, 2, figsize=(8.8, 7.2), constrained_layout=True)
    axes = axes.ravel()
    for ax, case in zip(axes, cases):
        im = draw_matrix_heatmap(ax, df, case, "reliability_raw", "viridis", vmin=0, vmax=1, title=CASE_LABELS[case])
        ax.set_xlabel("噪声标准差", fontsize=8)
        if ax in [axes[0], axes[2]]:
            ax.set_ylabel("观测点数", fontsize=8)
    for label, ax in zip("abcd", axes):
        add_panel_label(ax, label)
    cbar = fig.colorbar(im, ax=axes, location="right", shrink=0.84, pad=0.012)
    cbar.set_label("R_lin", fontsize=8)
    save_png(fig, FIG2_PATH)


def build_figure4() -> None:
    base = ROOT / "minimal_pinn" / "results" / "probability_matrices" / "burgers_probability_boundary_v2_5seed"
    df = pd.read_csv(base / "multiseed_summary_with_ci.csv")
    ci = pd.read_csv(base / "cross_rate_ci_by_noise.csv")
    fig = plt.figure(figsize=(12.6, 4.9), constrained_layout=True)
    gs = fig.add_gridspec(1, 3, width_ratios=[1.05, 1.05, 1.0], wspace=0.36)

    for idx, (value, title, cmap) in enumerate([
        ("crosses_threshold_rate", "越界率", "rocket_r" if False else "Reds"),
        ("cross_rate_ci_width", "Wilson 区间宽度", "YlGnBu"),
    ]):
        ax = fig.add_subplot(gs[0, idx])
        piv = df.pivot_table(index="num_observation", columns="noise_std", values=value, aggfunc="mean").sort_index(ascending=False).sort_index(axis=1)
        im = ax.imshow(piv.to_numpy(dtype=float), aspect="auto", cmap=cmap, vmin=0, vmax=1)
        ax.set_title(title, fontsize=9.8, fontweight="bold")
        ax.set_xticks(range(len(piv.columns)))
        ax.set_xticklabels([f"{x:.2f}" for x in piv.columns], rotation=45, ha="right", fontsize=7)
        ax.set_yticks(range(len(piv.index)))
        ax.set_yticklabels([str(int(x)) for x in piv.index], fontsize=7)
        ax.set_xlabel("噪声标准差", fontsize=8)
        ax.set_ylabel("观测点数" if idx == 0 else "", fontsize=8)
        ax.tick_params(length=0)
        for r in range(piv.shape[0]):
            for c in range(piv.shape[1]):
                val = piv.iloc[r, c]
                ax.text(c, r, f"{val:.1f}", ha="center", va="center", fontsize=7, color="#202020")
        add_panel_label(ax, "ab"[idx])
        cb = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.03)
        cb.ax.tick_params(labelsize=7)

    ax = fig.add_subplot(gs[0, 2])
    x = ci["noise_std"].to_numpy(dtype=float)
    y = ci["crosses_threshold_rate_mean"].to_numpy(dtype=float) if "crosses_threshold_rate_mean" in ci.columns else ci.iloc[:, 1].to_numpy(dtype=float)
    low_col = "cross_rate_ci_low"
    high_col = "cross_rate_ci_high"
    if low_col in ci.columns and high_col in ci.columns:
        lo = ci[low_col].to_numpy(dtype=float)
        hi = ci[high_col].to_numpy(dtype=float)
        ax.fill_between(x, lo, hi, color="#B44B4B", alpha=0.18, linewidth=0)
    ax.plot(x, y, marker="o", color="#B44B4B", lw=1.8)
    ax.set_ylim(-0.03, 1.03)
    ax.set_xlabel("噪声标准差", fontsize=8)
    ax.set_ylabel("平均越界率", fontsize=8)
    ax.set_title("随噪声扩张的统计边界", fontsize=9.8, fontweight="bold")
    ax.grid(alpha=0.25)
    add_panel_label(ax, "c")
    save_png(fig, FIG4_PATH)


def build_figure5() -> None:
    tables = []
    for case in ["poisson", "stokes_poiseuille", "fisher_kpp", "burgers"]:
        file_case = "stokes" if case == "stokes_poiseuille" else case
        path = ROOT / "minimal_pinn" / "results" / "analysis" / "indicator_validity_v1" / f"{file_case}_indicator_table.csv"
        d = pd.read_csv(path)
        d["case"] = case
        tables.append(d)
    df = pd.concat(tables, ignore_index=True)
    dims = list(DIMENSION_LABELS)
    degradation = df.groupby("case")[dims].mean().reindex(CASE_LABELS).pipe(lambda x: 1 - x)
    summary = pd.read_csv(ROOT / "minimal_pinn" / "results" / "analysis" / "single_vs_multi_v1" / "single_vs_multi_summary.csv").set_index("case")
    r2_cols = ["r2_physics_vs_rel_l2", "r2_training_vs_rel_l2", "r2_numerical_vs_rel_l2", "r2_structural_vs_rel_l2"]
    r2 = summary.reindex(CASE_LABELS)[r2_cols]

    fig = plt.figure(figsize=(13.8, 5.7), constrained_layout=True)
    gs = fig.add_gridspec(1, 3, width_ratios=[1.20, 1.05, 1.02], wspace=0.42)

    ax = fig.add_subplot(gs[0, 0])
    bottom = np.zeros(len(degradation))
    x = np.arange(len(degradation))
    dim_colors = ["#6B8EAD", "#D79B5A", "#7AA874", "#B66A6A"]
    for color, dim in zip(dim_colors, dims):
        vals = degradation[dim].to_numpy(dtype=float)
        ax.bar(x, vals, bottom=bottom, color=color, width=0.68, label=DIMENSION_LABELS[dim])
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels([CASE_LABELS[c] for c in degradation.index], rotation=18, ha="right", fontsize=8)
    ax.set_ylabel("平均退化强度（1 - 维度得分）", fontsize=8)
    ax.set_title("四维失效来源", fontsize=10, fontweight="bold")
    ax.legend(fontsize=7, frameon=False, ncol=2, loc="upper left")
    ax.grid(axis="y", alpha=0.22)
    add_panel_label(ax, "a")

    ax = fig.add_subplot(gs[0, 1])
    im = ax.imshow(r2.to_numpy(dtype=float), cmap="Blues", vmin=0, vmax=1, aspect="auto")
    ax.set_xticks(range(4))
    ax.set_xticklabels([DIMENSION_LABELS[d] for d in dims], rotation=35, ha="right", fontsize=7.5)
    ax.set_yticks(range(len(r2.index)))
    ax.set_yticklabels([CASE_LABELS[c] for c in r2.index], fontsize=8)
    for i in range(r2.shape[0]):
        for j in range(r2.shape[1]):
            ax.text(j, i, f"{r2.iloc[i, j]:.2f}", ha="center", va="center", fontsize=7, color="#1F2933")
    ax.set_title("相对 L2 对各维度的解释度", fontsize=10, fontweight="bold")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.03).set_label("$R^2$", fontsize=8)
    add_panel_label(ax, "b")

    ax = fig.add_subplot(gs[0, 2])
    dominant = degradation.idxmax(axis=1)
    y = np.arange(len(dominant))
    vals = degradation.max(axis=1).to_numpy(dtype=float)
    colors = [dim_colors[dims.index(d)] for d in dominant]
    ax.barh(y, vals, color=colors, height=0.55)
    for yy, val, dim in zip(y, vals, dominant):
        ax.text(val + 0.01, yy, DIMENSION_LABELS[dim], va="center", fontsize=8)
    ax.set_yticks(y)
    ax.set_yticklabels([CASE_LABELS[c] for c in dominant.index], fontsize=8)
    ax.set_xlim(0, max(vals) * 1.35)
    ax.set_xlabel("主导退化强度", fontsize=8)
    ax.set_title("主导失效维度", fontsize=10, fontweight="bold")
    ax.grid(axis="x", alpha=0.22)
    add_panel_label(ax, "c")
    save_png(fig, FIG5_PATH)


def build_figure6() -> None:
    df = pd.read_csv(STRATEGY_CSV)
    strategies = ["naive_region_aware_v1", "dim_guided_v2", "non_dominant_guided_v3"]
    labels = {
        "naive_region_aware_v1": "朴素采样",
        "dim_guided_v2": "主导维度",
        "non_dominant_guided_v3": "非主导维度",
    }
    colors = {"naive_region_aware_v1": "#D79B5A", "dim_guided_v2": "#7AA874", "non_dominant_guided_v3": "#B66A6A"}
    rows = []
    for (case, label), sub in df.groupby(["case", "label"]):
        base = sub[sub["strategy"] == "baseline"].iloc[0]
        for strategy in strategies:
            s = sub[sub["strategy"] == strategy].iloc[0]
            rows.append({
                "case": case,
                "label": label,
                "strategy": strategy,
                "delta_rel_l2": s["rel_l2_mean"] - base["rel_l2_mean"],
                "delta_R": s["reliability_raw_recal_mean"] - base["reliability_raw_recal_mean"],
                "rel_l2_std": s["rel_l2_std"],
                "R_std": s["reliability_raw_recal_std"],
            })
    out = pd.DataFrame(rows)
    fig, axes = plt.subplots(1, 2, figsize=(12.8, 4.9), constrained_layout=True)
    metrics = [("delta_rel_l2", "相对 L2 误差变化\n负值表示改进"), ("delta_R", "R_lin 变化\n正值表示改进")]
    for ax, (metric, title) in zip(axes, metrics):
        y_positions = []
        y_labels = []
        vals = []
        cols = []
        cur = 0
        for case in ["burgers", "stokes_poiseuille"]:
            case_rows = out[out["case"] == case]
            for _, r in case_rows.iterrows():
                y_positions.append(cur)
                prefix = "Burgers" if case == "burgers" else "Stokes"
                y_labels.append(f"{prefix} | {labels[r['strategy']]}")
                vals.append(r[metric])
                cols.append(colors[r["strategy"]])
                cur += 1
            cur += 0.8
        ax.axvline(0, color="#333333", lw=0.8)
        ax.barh(y_positions, vals, color=cols, height=0.55)
        ax.set_yticks(y_positions)
        ax.set_yticklabels(y_labels, fontsize=7.8)
        ax.set_title(title, fontsize=10, fontweight="bold")
        ax.grid(axis="x", alpha=0.22)
        add_panel_label(ax, "ab"[0 if metric == "delta_rel_l2" else 1])
    save_png(fig, FIG6_PATH)


def build_figure7() -> None:
    rows = [
        ("Poisson", "椭圆平滑\n全局正则性", "扰动较易被\n解结构吸收", "未形成\n实用边界", "#2F7D5C"),
        ("Stokes-\nPoiseuille", "线性耦合流动\n强边界约束", "观测不足时\n出现阈值", "窄硬边界\n参考解误差主导", "#3A6EA5"),
        ("Fisher-KPP", "反应-扩散前沿\n扩散平滑", "前沿位置\n逐步漂移", "规则边界带\n中等统计宽度", "#C18B21"),
        ("Burgers", "非线性输运\n陡峭梯度", "局部吸引区\n种子敏感", "宽异质临界带\n局部异常", "#B44B4B"),
    ]
    fig, ax = plt.subplots(figsize=(13.4, 6.0), constrained_layout=True)
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    headers = [("PDE", 0.08), ("算子性质", 0.30), ("传播与优化", 0.57), ("边界形态", 0.84)]
    for text, x in headers:
        ax.text(x, 0.93, text, ha="center", va="center", fontsize=10.2, fontweight="bold")
    y_positions = [0.77, 0.58, 0.39, 0.20]
    for y, (case, op, mechanism, topo, color) in zip(y_positions, rows):
        ax.text(0.08, y, case, ha="center", va="center", fontsize=9.8, color="white",
                bbox=dict(boxstyle="round,pad=0.42,rounding_size=0.04", facecolor=color, edgecolor=color))
        for x, text, width in [(0.30, op, 0.19), (0.57, mechanism, 0.23), (0.84, topo, 0.21)]:
            ax.text(x, y, text, ha="center", va="center", fontsize=8.9, linespacing=1.25,
                    bbox=dict(boxstyle="round,pad=0.42,rounding_size=0.04", facecolor="#F8F8F6", edgecolor=color, linewidth=1.1))
        ax.annotate("", xy=(0.43, y), xytext=(0.39, y), arrowprops=dict(arrowstyle="->", lw=1.3, color="#444444"))
        ax.annotate("", xy=(0.71, y), xytext=(0.67, y), arrowprops=dict(arrowstyle="->", lw=1.3, color="#444444"))
    ax.text(0.50, 0.055, "经验解释：四维诊断状态用于分析失效来源，标量分数只用于排序、绘图和边界定位。", ha="center", fontsize=8.6, color="#333333")
    save_png(fig, FIG7_PATH)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("\\times", "×")
    text = text.replace("\\ge", "≥")
    text = text.replace("\\le", "≤")
    text = text.replace("\\sigma", "σ")
    text = text.replace("\\pm", "±")
    text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\mathcal\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\left|\\right", "", text)
    text = text.replace("\\", "")
    return text.strip()


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.12
    style.paragraph_format.space_after = Pt(4)

    for name, size in [("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10.5), ("Heading 4", 10.5)]:
        if name in doc.styles:
            h = doc.styles[name]
            h.font.name = "Times New Roman"
            h._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            h.font.size = Pt(size)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(5)


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None, bold=False, font_size: float = 10.5) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    elif style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(clean_inline(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {min(level, 4)}"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(clean_inline(text))
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(12 if level == 1 else 11)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        if edge_data:
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
        else:
            element.set(qn("w:val"), "nil")


def style_table_text(table, font_size: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(font_size)


def apply_three_line_table(table) -> None:
    line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    thin = {"val": "single", "sz": "4", "space": "0", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=line, bottom=thin, left=None, right=None, insideH=None, insideV=None)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=line, left=None, right=None, insideH=None, insideV=None)


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in lines:
        if re.match(r"^\|\s*-", raw):
            continue
        cells = [clean_inline(cell.strip()) for cell in raw.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = parse_markdown_table(table_lines)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            if r_idx == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
            if re.fullmatch(r"[-+]?\d+(\.\d+)?", value):
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_size = 8.0 if len(rows[0]) >= 6 else 8.8
    style_table_text(table, font_size)
    apply_three_line_table(table)
    doc.add_paragraph()


def insert_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        add_paragraph(doc, f"[缺失图像] {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.15))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(clean_inline(caption))
    cr.italic = False
    cr.font.name = "Times New Roman"
    cr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cr.font.size = Pt(9)
    cp.paragraph_format.space_after = Pt(6)


def build_doc() -> Document:
    build_figure1()
    build_figure2()
    build_figure4()
    build_figure5()
    build_figure6()
    build_figure7()
    text = MAIN_MD.read_text(encoding="utf-8").splitlines()
    fig_map: dict[str, list[dict]] = {}
    for fig in FIGURES:
        fig_map.setdefault(fig["section"], []).append(fig)

    doc = Document()
    set_doc_defaults(doc)

    title_added = False
    current_section = None
    inserted_for_section: set[str] = set()

    def insert_pending_figures() -> None:
        nonlocal current_section
        if current_section and current_section in fig_map and current_section not in inserted_for_section:
            for fig in fig_map[current_section]:
                insert_figure(doc, fig["path"], fig["caption"])
            inserted_for_section.add(current_section)

    i = 0
    while i < len(text):
        raw = text[i]
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(text) and text[i].rstrip().startswith("|"):
                table_lines.append(text[i].rstrip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if line.startswith("# "):
            insert_pending_figures()
            if not title_added:
                add_paragraph(doc, line[2:], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16)
                title_added = True
            else:
                add_heading(doc, line[2:], 1)
            i += 1
            continue
        if line.startswith("## "):
            insert_pending_figures()
            add_paragraph(doc, line[3:], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=12)
            i += 1
            continue
        if line.startswith("### "):
            insert_pending_figures()
            current_section = line
            add_heading(doc, line[4:], 1)
            i += 1
            continue
        if line.startswith("#### "):
            add_heading(doc, line[5:], 2)
            i += 1
            continue
        if re.match(r"^表\s*\d+[.．]", line):
            add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9)
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            add_paragraph(doc, line)
            i += 1
            continue
        if line.startswith("- "):
            add_paragraph(doc, "• " + line[2:])
            i += 1
            continue
        add_paragraph(doc, line)
        i += 1

    insert_pending_figures()

    return doc


def main() -> None:
    if DOCX_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)
    doc = build_doc()
    doc.save(DOCX_PATH)
    print(f"[done] backup={BACKUP_PATH}")
    print(f"[done] docx={DOCX_PATH}")


if __name__ == "__main__":
    main()
