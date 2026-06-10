"""
Generate the complete paper manuscript as a .docx file.
Reads markdown sections, creates figures from experimental data,
and assembles everything into a formatted document.
"""

import os
import re
import csv
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.colors import LogNorm
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from pathlib import Path

# ── Paths ──────────────────────────────────────────────────────
ROOT = Path(r"C:\Users\enmusubi4\Desktop\pinn-reliability-paper")
PAPER_DIR = ROOT / "paper"
RESULTS = ROOT / "minimal_pinn" / "results"
FIGURES_DIR = PAPER_DIR / "figures"
FIGURES_DIR.mkdir(exist_ok=True)

# ── Data Loading ───────────────────────────────────────────────
def load_csv(path):
    with open(path, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        return list(reader)

def get_matrix_data(name):
    path = RESULTS / "matrices" / name / "matrix_summary.csv"
    if path.exists():
        return load_csv(path)
    return []

# ── Plot Style ─────────────────────────────────────────────────
plt.rcParams.update({
    'font.size': 9,
    'axes.titlesize': 10,
    'axes.labelsize': 9,
    'xtick.labelsize': 8,
    'ytick.labelsize': 8,
    'legend.fontsize': 8,
    'figure.dpi': 300,
    'savefig.dpi': 300,
    'savefig.bbox': 'tight',
})


# ═══════════════════════════════════════════════════════════════
# FIGURE GENERATORS — one figure per PDE case
# ═══════════════════════════════════════════════════════════════

def make_single_heatmap(data, case_name, title, filename, vmin=None, vmax=None):
    """Generate a single heatmap figure for one PDE case."""
    obs_levels = sorted(set(int(r['num_observation']) for r in data), reverse=True)
    noise_levels = sorted(set(float(r['noise_std']) for r in data))

    matrix = np.full((len(noise_levels), len(obs_levels)), np.nan)
    for r in data:
        oi = obs_levels.index(int(r['num_observation']))
        ni = noise_levels.index(float(r['noise_std']))
        matrix[ni, oi] = float(r['rel_l2'])

    fig, ax = plt.subplots(figsize=(4.5, 3.5))
    im = ax.imshow(matrix, aspect='auto', cmap='YlOrRd', vmin=vmin, vmax=vmax,
                   origin='lower', interpolation='nearest')
    ax.set_xticks(range(len(obs_levels)))
    ax.set_xticklabels(obs_levels)
    ax.set_yticks(range(len(noise_levels)))
    ax.set_yticklabels([f'{s:.2f}' for s in noise_levels])
    ax.set_xlabel('Observations')
    ax.set_ylabel('Noise σ')
    ax.set_title(title)
    fig.colorbar(im, ax=ax, label='rel_l2', shrink=0.8)
    plt.tight_layout()
    path = FIGURES_DIR / filename
    fig.savefig(path)
    plt.close(fig)
    return path


def make_single_probability_heatmap(prob_path, case_name, title, filename):
    """Generate a single probability boundary heatmap for one PDE case."""
    if not prob_path.exists():
        return None
    data = load_csv(prob_path)
    case_data = [r for r in data if r['case'] == case_name]
    if not case_data:
        return None

    obs_levels = sorted(set(int(r['num_observation']) for r in case_data), reverse=True)
    noise_levels = sorted(set(float(r['noise_std']) for r in case_data))

    matrix = np.full((len(noise_levels), len(obs_levels)), np.nan)
    for r in case_data:
        oi = obs_levels.index(int(r['num_observation']))
        ni = noise_levels.index(float(r['noise_std']))
        matrix[ni, oi] = float(r['crosses_threshold_rate'])

    fig, ax = plt.subplots(figsize=(4.5, 3.5))
    im = ax.imshow(matrix, aspect='auto', cmap='YlOrRd', vmin=0, vmax=1,
                   origin='lower', interpolation='nearest')
    ax.set_xticks(range(len(obs_levels)))
    ax.set_xticklabels(obs_levels, rotation=45)
    ax.set_yticks(range(len(noise_levels)))
    ax.set_yticklabels([f'{s:.2f}' for s in noise_levels])
    ax.set_xlabel('Observations')
    ax.set_ylabel('Noise σ')
    ax.set_title(title)
    fig.colorbar(im, ax=ax, label='Crossing rate', shrink=0.8)
    plt.tight_layout()
    path = FIGURES_DIR / filename
    fig.savefig(path)
    plt.close(fig)
    return path


def make_protocol_b_curve(data_dict, title, filename):
    """Generate a single Protocol B degradation curve for one PDE."""
    obs_levels = sorted(set(k[0] for k in data_dict.keys()))
    noise_levels = sorted(set(k[1] for k in data_dict.keys()))
    colors = plt.cm.tab10(np.linspace(0, 1, len(noise_levels)))

    fig, ax = plt.subplots(figsize=(5, 4))
    for ni, noise in enumerate(noise_levels):
        rel_l2_vals = [data_dict.get((obs, noise), np.nan) for obs in obs_levels]
        ax.plot(obs_levels, rel_l2_vals, 'o-', label=f'σ={noise:.2f}',
                markersize=4, color=colors[ni])

    ax.set_xlabel('Observations')
    ax.set_ylabel('rel_l2')
    ax.set_title(title)
    ax.set_xscale('log')
    ax.set_yscale('log')
    ax.legend(fontsize=7, ncol=2)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    path = FIGURES_DIR / filename
    fig.savefig(path)
    plt.close(fig)
    return path


def make_loss_history_comparison():
    """Generate loss history comparison: Allen-Cahn vs Burgers."""
    ac_hist = RESULTS / "convergence_v1500" / "allen_cahn_e1500_obs128_noise000_seed42" / "history.csv"
    bu_hist = RESULTS / "convergence_v1500" / "burgers_e1500_obs128_noise000_seed42" / "history.csv"

    if not ac_hist.exists() or not bu_hist.exists():
        return None

    fig, ax = plt.subplots(figsize=(5, 4))
    for path, label, color in [(ac_hist, 'Allen-Cahn', 'C0'), (bu_hist, 'Burgers', 'C1')]:
        hist_data = load_csv(path)
        epochs = [int(float(r['epoch'])) for r in hist_data]
        losses = [float(r['loss_total']) for r in hist_data]
        ax.plot(epochs, losses, label=label, color=color, alpha=0.7)

    ax.set_xlabel('Epoch')
    ax.set_ylabel('Total loss')
    ax.set_title('Training loss: Allen-Cahn vs Burgers')
    ax.set_yscale('log')
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    path = FIGURES_DIR / "fig_loss_history.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def make_kdv_comparison():
    """Generate KdV single vs double soliton comparison."""
    single_obs = [16, 32, 64, 128, 256, 512]
    single_rel2 = [0.0225, 0.0100, 0.0057, 0.0054, 0.0042, 0.0046]
    double_obs = [16, 32, 64, 128, 256, 512]
    double_rel2 = [0.4394, 0.2014, 0.0992, 0.0194, 0.0200, 0.0134]

    fig, ax = plt.subplots(figsize=(5, 4))
    ax.plot(single_obs, single_rel2, 'o-', label='KdV Single (dim=1)', color='C0', markersize=5)
    ax.plot(double_obs, double_rel2, 's-', label='KdV Double (dim=2)', color='C1', markersize=5)
    ax.set_xlabel('Observations')
    ax.set_ylabel('rel_l2')
    ax.set_title('KdV: single vs double soliton (σ=0)')
    ax.set_xscale('log')
    ax.set_yscale('log')
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    path = FIGURES_DIR / "fig_kdv_comparison.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def make_gradient_bar():
    """Generate quantitative gradient bar chart."""
    categories = ['Stokes', 'Allen-\nCahn', 'Fisher-\nKPP', 'KdV\nSingle', 'NLS', 'Wave', 'Burgers', 'KdV\nDouble']

    baselines = [0.010, 0.008, 0.013, 0.004, 0.009, 0.008, 0.018, 0.013]
    seed_stds = [0.003, 0.002, 0.004, 0.001, 0.003, 0.008, 0.010, 0.005]
    deg_ratios = [3.9, 12, 5.6, 34, 28, 34, 5.4, 39]

    fig, axes = plt.subplots(1, 3, figsize=(10, 3.5))

    axes[0].bar(range(len(categories)), baselines, color='steelblue')
    axes[0].set_xticks(range(len(categories)))
    axes[0].set_xticklabels(categories, fontsize=7)
    axes[0].set_ylabel('Clean baseline rel_l2')
    axes[0].set_title('(a) Baseline accuracy')

    axes[1].bar(range(len(categories)), seed_stds, color='coral')
    axes[1].set_xticks(range(len(categories)))
    axes[1].set_xticklabels(categories, fontsize=7)
    axes[1].set_ylabel('Seed std (clean)')
    axes[1].set_title('(b) Seed variance')

    colors = ['green' if d < 10 else 'orange' if d < 30 else 'red' for d in deg_ratios]
    axes[2].bar(range(len(categories)), deg_ratios, color=colors)
    axes[2].set_xticks(range(len(categories)))
    axes[2].set_xticklabels(categories, fontsize=7)
    axes[2].set_ylabel('Degradation ratio')
    axes[2].set_title('(c) Degradation magnitude')

    fig.suptitle('Figure. Quantitative gradient across degradation archetypes', fontsize=11, y=1.02)
    plt.tight_layout()
    path = FIGURES_DIR / "fig_gradient.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def make_concept_diagram():
    """Generate the four-factor concept diagram."""
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis('off')

    props = [
        (1, 4.2, 'PDE Operator\nSpectral Properties'),
        (1, 3.2, 'Nonlinearity\nType'),
        (1, 2.2, 'Solution\nStructure'),
        (1, 1.2, 'Domain &\nBoundary'),
    ]
    for x, y, text in props:
        ax.add_patch(plt.Rectangle((x-0.4, y-0.3), 1.8, 0.6, fill=True,
                                    facecolor='lightblue', edgecolor='steelblue', linewidth=1.5))
        ax.text(x+0.5, y, text, ha='center', va='center', fontsize=8, fontweight='bold')

    factors = [
        (4.5, 4.2, 'Null Space\nDimension', '→ deg. directions'),
        (4.5, 3.2, 'Loss\nCurvature', '→ transition width'),
        (4.5, 2.2, 'Optimization\nMultimodality', '→ seed variance'),
        (4.5, 1.2, 'Information\nDensity', '→ regularity'),
    ]
    for x, y, text, sub in factors:
        ax.add_patch(plt.Rectangle((x-0.5, y-0.3), 2.0, 0.6, fill=True,
                                    facecolor='lightyellow', edgecolor='orange', linewidth=1.5))
        ax.text(x+0.5, y+0.05, text, ha='center', va='center', fontsize=8, fontweight='bold')
        ax.text(x+0.5, y-0.2, sub, ha='center', va='center', fontsize=6, style='italic')

    behaviors = [
        (8, 4.2, 'No degradation\n(Poisson)'),
        (8, 3.2, 'Narrow boundary\n(Stokes, Allen-Cahn)'),
        (8, 2.2, 'Wide probabilistic\nband (Burgers)'),
        (8, 1.2, 'Irregular/patchy\n(Heat equation)'),
    ]
    for x, y, text in behaviors:
        ax.add_patch(plt.Rectangle((x-0.5, y-0.3), 2.0, 0.6, fill=True,
                                    facecolor='lightgreen', edgecolor='green', linewidth=1.5))
        ax.text(x+0.5, y, text, ha='center', va='center', fontsize=8, fontweight='bold')

    for y in [4.2, 3.2, 2.2, 1.2]:
        ax.annotate('', xy=(4.0, y), xytext=(2.3, y),
                    arrowprops=dict(arrowstyle='->', color='gray', lw=1.5))
        ax.annotate('', xy=(7.0, y), xytext=(6.5, y),
                    arrowprops=dict(arrowstyle='->', color='gray', lw=1.5))

    ax.set_title('Four-factor degradation mechanism model', fontsize=11, pad=10)
    path = FIGURES_DIR / "fig_concept.png"
    fig.savefig(path)
    plt.close(fig)
    return path


def make_multidim_figures():
    """Generate multidimensional analysis figures."""
    paths = {}

    # (a) Dominant dimension distribution
    fig, ax = plt.subplots(figsize=(5, 4))
    cases = ['Stokes', 'Fisher-KPP', 'Burgers']
    physics = [1, 2, 9]
    training = [0, 5, 23]
    numerical = [7, 5, 10]
    structural = [0, 1, 17]
    x = np.arange(len(cases))
    width = 0.2
    ax.bar(x - 1.5*width, physics, width, label='Physics', color='C0')
    ax.bar(x - 0.5*width, training, width, label='Training', color='C1')
    ax.bar(x + 0.5*width, numerical, width, label='Numerical', color='C2')
    ax.bar(x + 1.5*width, structural, width, label='Structural', color='C3')
    ax.set_xticks(x)
    ax.set_xticklabels(cases)
    ax.set_ylabel('Count')
    ax.set_title('Dominant dimension distribution')
    ax.legend(fontsize=7)
    plt.tight_layout()
    p = FIGURES_DIR / "fig_dominant_dim.png"
    fig.savefig(p)
    plt.close(fig)
    paths['dominant'] = p

    # (b) training_stability vs rel_l2
    fig, ax = plt.subplots(figsize=(5, 4))
    coarse_data = get_matrix_data('coarse_v2')
    burgers_data = [r for r in coarse_data if r['case'] == 'burgers']
    if burgers_data:
        rel_l2_vals = [float(r['rel_l2']) for r in burgers_data]
        loss_std_vals = [float(r['loss_std']) for r in burgers_data]
        ax.scatter(rel_l2_vals, loss_std_vals, alpha=0.5, s=20, color='C1')
        ax.set_xlabel('rel_l2')
        ax.set_ylabel('loss_std')
        x_arr = np.array(rel_l2_vals)
        y_arr = np.array(loss_std_vals)
        if len(x_arr) > 2:
            corr = np.corrcoef(x_arr, y_arr)[0, 1]
            ax.set_title(f'Burgers: training_stability vs rel_l2 (R²={corr**2:.3f})')
        ax.grid(True, alpha=0.3)
    plt.tight_layout()
    p = FIGURES_DIR / "fig_training_vs_error.png"
    fig.savefig(p)
    plt.close(fig)
    paths['training'] = p

    # (c) Dimension ablation
    fig, ax = plt.subplots(figsize=(5, 4))
    ablation_cases = ['Stokes', 'Fisher-KPP', 'Burgers']
    full_4d = [0.95, 0.93, 0.75]
    rel_l2_only = [0.93, 0.92, 0.60]
    x = np.arange(len(ablation_cases))
    ax.bar(x - 0.2, full_4d, 0.35, label='Full 4D R', color='steelblue')
    ax.bar(x + 0.2, rel_l2_only, 0.35, label='rel_l2 only', color='coral')
    ax.set_xticks(x)
    ax.set_xticklabels(ablation_cases)
    ax.set_ylabel('Cross-seed Spearman ρ')
    ax.set_title('Dimension ablation')
    ax.legend(fontsize=7)
    ax.set_ylim(0.5, 1.0)
    plt.tight_layout()
    p = FIGURES_DIR / "fig_ablation.png"
    fig.savefig(p)
    plt.close(fig)
    paths['ablation'] = p

    return paths


# ═══════════════════════════════════════════════════════════════
# MAIN: generate all figures
# ═══════════════════════════════════════════════════════════════

def generate_all_figures():
    """Generate all individual figures and return a dict of paths."""
    fig_paths = {}

    # ── Phase maps (one per case) ──
    phase_cases = [
        ('coarse_v2', 'poisson', 'Poisson: rel_l2', 'fig01_poisson_phase.png'),
        ('coarse_v2', 'stokes_poiseuille', 'Stokes-Poiseuille: rel_l2', 'fig02_stokes_phase.png'),
        ('coarse_v2_allen_cahn', 'allen_cahn', 'Allen-Cahn: rel_l2', 'fig03_allen_cahn_phase.png'),
        ('coarse_v2', 'fisher_kpp', 'Fisher-KPP: rel_l2', 'fig04_fisher_kpp_phase.png'),
        ('coarse_v2', 'burgers', 'Burgers: rel_l2', 'fig05_burgers_phase.png'),
        ('coarse_v2_heat', 'heat_equation', 'Heat Equation: rel_l2', 'fig06_heat_phase.png'),
    ]
    for matrix_name, case_name, title, filename in phase_cases:
        data = get_matrix_data(matrix_name)
        case_data = [r for r in data if r['case'] == case_name]
        if case_data:
            fig_paths[filename] = make_single_heatmap(case_data, case_name, title, filename)

    # ── Probability boundaries (one per case) ──
    prob_cases = [
        ('probability_v2_stokes', 'stokes_poiseuille', 'Stokes-Poiseuille: crossing rate', 'fig07_stokes_prob.png'),
        ('probability_v2_fisher_kpp', 'fisher_kpp', 'Fisher-KPP: crossing rate', 'fig08_fisher_kpp_prob.png'),
        ('probability_v2_burgers', 'burgers', 'Burgers: crossing rate', 'fig09_burgers_prob.png'),
    ]
    for dir_name, case_name, title, filename in prob_cases:
        prob_path = RESULTS / "probability_matrices" / dir_name / "multiseed_summary.csv"
        p = make_single_probability_heatmap(prob_path, case_name, title, filename)
        if p:
            fig_paths[filename] = p

    # ── Protocol B degradation curves (one per case) ──
    # KdV single soliton
    kdv_data = {
        (512, 0.0): 0.0046, (512, 0.05): 0.0077, (512, 0.10): 0.0132, (512, 0.15): 0.0192, (512, 0.20): 0.0254,
        (256, 0.0): 0.0042, (256, 0.05): 0.0114, (256, 0.10): 0.0215, (256, 0.15): 0.0321, (256, 0.20): 0.0430,
        (128, 0.0): 0.0054, (128, 0.05): 0.0135, (128, 0.10): 0.0256, (128, 0.15): 0.0384, (128, 0.20): 0.0515,
        (64, 0.0): 0.0057, (64, 0.05): 0.0224, (64, 0.10): 0.0440, (64, 0.15): 0.0671, (64, 0.20): 0.0916,
        (32, 0.0): 0.0100, (32, 0.05): 0.0335, (32, 0.10): 0.0646, (32, 0.15): 0.0966, (32, 0.20): 0.1292,
        (16, 0.0): 0.0225, (16, 0.05): 0.0397, (16, 0.10): 0.0697, (16, 0.15): 0.1015, (16, 0.20): 0.1349,
    }
    fig_paths['fig10_kdv_curve.png'] = make_protocol_b_curve(kdv_data, 'KdV Single Soliton', 'fig10_kdv_curve.png')

    # NLS soliton
    nls_data = {
        (512, 0.0): 0.0078, (512, 0.05): 0.0111, (512, 0.10): 0.0178, (512, 0.15): 0.0253, (512, 0.20): 0.0331,
        (256, 0.0): 0.0093, (256, 0.05): 0.0152, (256, 0.10): 0.0258, (256, 0.15): 0.0373, (256, 0.20): 0.0493,
        (128, 0.0): 0.0096, (128, 0.05): 0.0176, (128, 0.10): 0.0321, (128, 0.15): 0.0479, (128, 0.20): 0.0643,
        (64, 0.0): 0.0161, (64, 0.05): 0.0275, (64, 0.10): 0.0477, (64, 0.15): 0.0699, (64, 0.20): 0.0932,
        (32, 0.0): 0.0390, (32, 0.05): 0.0570, (32, 0.10): 0.0838, (32, 0.15): 0.1125, (32, 0.20): 0.1426,
        (16, 0.0): 0.0216, (16, 0.05): 0.0455, (16, 0.10): 0.0843, (16, 0.15): 0.1248, (16, 0.20): 0.1637,
    }
    fig_paths['fig11_nls_curve.png'] = make_protocol_b_curve(nls_data, 'NLS Soliton', 'fig11_nls_curve.png')

    # Wave equation
    wave_data = {
        (256, 0.0): 0.0081, (256, 0.05): 0.0104, (256, 0.10): 0.0151, (256, 0.15): 0.0209, (256, 0.20): 0.0270,
        (128, 0.0): 0.0074, (128, 0.05): 0.0124, (128, 0.10): 0.0200, (128, 0.15): 0.0286, (128, 0.20): 0.0369,
        (64, 0.0): 0.0094, (64, 0.05): 0.0157, (64, 0.10): 0.0276, (64, 0.15): 0.0409, (64, 0.20): 0.0541,
        (32, 0.0): 0.0099, (32, 0.05): 0.0192, (32, 0.10): 0.0349, (32, 0.15): 0.0514, (32, 0.20): 0.0682,
        (16, 0.0): 0.0122, (16, 0.05): 0.0386, (16, 0.10): 0.0752, (16, 0.15): 0.1127, (16, 0.20): 0.1512,
        (8, 0.0): 0.0172, (8, 0.05): 0.0338, (8, 0.10): 0.0578, (8, 0.15): 0.0813, (8, 0.20): 0.1054,
    }
    fig_paths['fig12_wave_curve.png'] = make_protocol_b_curve(wave_data, 'Wave Equation (1st order)', 'fig12_wave_curve.png')

    # ── Comparison figures ──
    fig_paths['fig13_kdv_comparison.png'] = make_kdv_comparison()
    fig_paths['fig_loss_history.png'] = make_loss_history_comparison()

    # ── Mechanism figures ──
    fig_paths['fig_concept.png'] = make_concept_diagram()

    # ── Gradient bar chart ──
    fig_paths['fig_gradient.png'] = make_gradient_bar()

    # ── Multidimensional figures ──
    md_paths = make_multidim_figures()
    fig_paths.update(md_paths)

    return fig_paths


# ═══════════════════════════════════════════════════════════════
# DOCUMENT ASSEMBLY
# ═══════════════════════════════════════════════════════════════

def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
        elif level == 2:
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True
        else:
            heading_style.font.size = Pt(11)
            heading_style.font.bold = True


def add_figure(doc, image_path, caption):
    if image_path and Path(image_path).exists():
        doc.add_picture(str(image_path), width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


def add_table_from_data(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

    n_cols = len(headers)
    clean_rows = []
    for row in rows:
        if len(row) >= n_cols:
            clean_rows.append(row[:n_cols])
        else:
            clean_rows.append(row + [''] * (n_cols - len(row)))

    table = doc.add_table(rows=len(clean_rows)+1, cols=n_cols)
    table.style = 'Light Grid Accent 1'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for r_idx, row in enumerate(clean_rows):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(row[c_idx])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()


def read_markdown_section(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()


def add_markdown_content(doc, content):
    lines = content.split('\n')
    in_table = False
    table_rows = []
    table_headers = []
    in_code_block = False
    code_content = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line)
            continue

        if not stripped:
            if in_table and table_headers:
                add_table_from_data(doc, table_headers, table_rows)
                table_rows = []
                table_headers = []
                in_table = False
            continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            if in_table and table_headers:
                add_table_from_data(doc, table_headers, table_rows)
                table_rows = []
                table_headers = []
                in_table = False
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        elif stripped.startswith('## '):
            if in_table and table_headers:
                add_table_from_data(doc, table_headers, table_rows)
                table_rows = []
                table_headers = []
                in_table = False
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        elif stripped.startswith('### '):
            if in_table and table_headers:
                add_table_from_data(doc, table_headers, table_rows)
                table_rows = []
                table_headers = []
                in_table = False
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(re.match(r'^[-:\s]+$', c) for c in cells if c):
                continue
            if all(c == '' for c in cells):
                continue
            if not in_table:
                table_headers = cells
                in_table = True
            else:
                table_rows.append(cells)
            continue

        if in_table and table_headers:
            add_table_from_data(doc, table_headers, table_rows)
            table_rows = []
            table_headers = []
            in_table = False

        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)

    if in_table and table_headers:
        add_table_from_data(doc, table_headers, table_rows)


def build_document():
    print("Generating figures...")
    fig_paths = generate_all_figures()

    print("Creating document...")
    doc = Document()
    setup_styles(doc)

    # Title
    title = doc.add_heading('物理信息神经网络在稀疏与噪声观测下的可靠性退化机制：跨偏微分方程系统的受控扫描研究', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    doc.add_heading('摘要', level=1)
    add_markdown_content(doc, read_markdown_section(PAPER_DIR / '00_title_abstract.md').split('\n', 3)[-1])

    # Sections
    sections = [
        ('01_introduction.md', None),
        ('02_related_work.md', None),
        ('03_method.md', None),
        ('04_experiment_design.md', None),
        ('05_results_archetypes.md', None),
        ('06_results_mechanism.md', None),
        ('07_results_multidimensional.md', None),
        ('08_discussion.md', None),
        ('09_conclusion.md', None),
        ('10_appendix.md', None),
    ]

    for filename, _ in sections:
        filepath = PAPER_DIR / filename
        if filepath.exists():
            doc.add_page_break()
            add_markdown_content(doc, read_markdown_section(filepath))

    # Insert figures at appropriate positions
    # Note: In the final document, figures should be referenced by the markdown text.
    # For now, we add a figure appendix at the end.

    doc.add_page_break()
    doc.add_heading('图版', level=1)

    # Phase maps
    phase_figs = [
        ('fig01_poisson_phase.png', '图1. 泊松方程相对L2误差相空间。'),
        ('fig02_stokes_phase.png', '图2. 斯托克斯-泊肃叶流相对L2误差相空间。'),
        ('fig03_allen_cahn_phase.png', '图3. Allen-Cahn方程相对L2误差相空间。'),
        ('fig04_fisher_kpp_phase.png', '图4. Fisher-KPP方程相对L2误差相空间。'),
        ('fig05_burgers_phase.png', '图5. Burgers方程相对L2误差相空间。'),
        ('fig06_heat_phase.png', '图6. 热方程相对L2误差相空间。'),
    ]
    for fname, cap in phase_figs:
        add_figure(doc, fig_paths.get(fname), cap)

    # Probability boundaries
    prob_figs = [
        ('fig07_stokes_prob.png', '图7. 斯托克斯-泊肃叶流概率边界。'),
        ('fig08_fisher_kpp_prob.png', '图8. Fisher-KPP概率边界。'),
        ('fig09_burgers_prob.png', '图9. Burgers概率边界。'),
    ]
    for fname, cap in prob_figs:
        add_figure(doc, fig_paths.get(fname), cap)

    # Protocol B curves
    pb_figs = [
        ('fig10_kdv_curve.png', '图10. KdV单孤子退化曲线。'),
        ('fig11_nls_curve.png', '图11. NLS孤子退化曲线。'),
        ('fig12_wave_curve.png', '图12. 波方程退化曲线。'),
    ]
    for fname, cap in pb_figs:
        add_figure(doc, fig_paths.get(fname), cap)

    # Comparison and mechanism
    other_figs = [
        ('fig13_kdv_comparison.png', '图13. KdV单孤子vs双孤子退化对比。'),
        ('fig_loss_history.png', '图14. Allen-Cahn vs Burgers训练loss历史。'),
        ('fig_gradient.png', '图15. 三种退化原型定量梯度。'),
        ('fig_concept.png', '图16. 四因素退化机制概念图。'),
        ('fig_dominant_dim.png', '图17. 主导维度分布。'),
        ('fig_training_vs_error.png', '图18. Burgers训练稳定性vs rel_l2。'),
        ('fig_ablation.png', '图19. 维度消融分析。'),
    ]
    for fname, cap in other_figs:
        add_figure(doc, fig_paths.get(fname), cap)

    # Tables
    doc.add_page_break()
    doc.add_heading('表版', level=1)

    add_table_from_data(doc,
        ['Case', 'PDE Type', 'Null Dim', 'Protocol', 'Clean rel_l2', 'Archetype'],
        [
            ['Poisson', 'Elliptic, linear', '0', 'A', '0.097', 'No degradation'],
            ['Stokes-Poiseuille', 'Saddle-point, linear', '1', 'A', '0.010', 'Narrow boundary'],
            ['Allen-Cahn', 'Elliptic, nonlinear', '1', 'A', '0.008', 'Narrow boundary'],
            ['Fisher-KPP', 'Parabolic, weak NL', '1', 'A', '0.013', 'Medium boundary'],
            ['Burgers', 'Parabolic, strong NL', '3+', 'A', '0.018', 'Wide probabilistic'],
            ['Heat equation', 'Parabolic, linear', '1', 'A', '0.016', 'Narrow, irregular'],
            ['KdV single', 'Dispersive, NL', '1', 'B', '0.004', 'Narrow boundary'],
            ['NLS soliton', 'Dispersive, NL', '1', 'B', '0.009', 'Narrow boundary'],
            ['Wave equation', 'Hyperbolic, linear', '1', 'B', '0.008', 'Narrow boundary'],
            ['KdV double', 'Dispersive, NL', '2', 'B', '0.013', 'Wider boundary'],
        ],
        caption='表1. 十个PDE案例概览。')

    # Save
    output_path = ROOT / "paper_manuscript.docx"
    doc.save(str(output_path))
    print(f"\nDocument saved to: {output_path}")
    print(f"Figures saved to: {FIGURES_DIR}")
    return output_path


if __name__ == '__main__':
    build_document()
