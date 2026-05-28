"""
Generate formatted .docx from paper_manuscript_zh.md
- Chinese academic formatting (三线表, proper fonts, spacing)
- Embedded figures at correct positions
- Three-line tables
"""

from __future__ import annotations

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT / "paper_manuscript_zh.md"
FIG_DIR = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v3"
MORPH_DIR = PROJECT / "minimal_pinn" / "results" / "analysis" / "divergence_morphology_v1"
OUTPUT = PROJECT / "paper_manuscript_zh.docx"

# Figure mapping with subfigure labels
FIGURES = {
    "图 1-(a)": FIG_DIR / "fig01_rel_l2_ab.png",
    "图 1-(b)": FIG_DIR / "fig01_rel_l2_cd.png",
    "图 2-(a)": FIG_DIR / "fig02_R_ab.png",
    "图 2-(b)": FIG_DIR / "fig02_R_cd.png",
    "图 3-(a)": FIG_DIR / "fig03_boundary_a.png",
    "图 3-(b)": FIG_DIR / "fig03_boundary_b.png",
    "图 3-(c)": FIG_DIR / "fig03_boundary_c.png",
    "图 4-(a)": FIG_DIR / "fig04_ablation_a.png",
    "图 4-(b)": FIG_DIR / "fig04_ablation_b.png",
    "图 4-(c)": FIG_DIR / "fig04_ablation_c.png",
    "图 5": MORPH_DIR / "burgers_R-worst_not_L2-worst.png",
    "图 6": MORPH_DIR / "burgers_L2-worst_not_R-worst.png",
    "图 A1": PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v2" / "fig_a1_calibration_sensitivity.png",
    "图 A2": PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v2" / "fig_a2_anti_circularity.png",
    "图 A3": PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v2" / "fig_a3_baseline_failure.png",
}


def add_paragraph(doc, text, style="Normal", bold=False, font_size=10.5, alignment=None, spacing_after=6):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if bold:
        run.bold = True
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    pf.line_spacing = 1.5
    return p


def add_heading_custom(doc, text, level=1):
    """Add heading with proper Chinese font."""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 0:
        run.font.size = Pt(16)
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    elif level == 3:
        run.font.size = Pt(10.5)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True
    return h


def format_three_line_table(table):
    """Apply three-line-table style: top border, header bottom border, bottom border only."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    
    # Remove all existing borders first
    for cell in table._cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="nil"/>')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Set table borders
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:left w:val="nil"/>'
        f'<w:right w:val="nil"/>'
        f'<w:insideH w:val="nil"/>'
        f'<w:insideV w:val="nil"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # Bottom border for header row (first row)
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)
        bottom = parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="6" w:space="0" w:color="000000"/>')
        tcBorders.append(bottom)

    # Set table width
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def add_figure(doc, fig_path, caption, width_inches=6.0):
    """Add a figure with caption."""
    if not fig_path.exists():
        p = add_paragraph(doc, f"[Figure not found: {fig_path.name}]", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return
    
    # Add caption above figure
    add_paragraph(doc, caption, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=4)
    
    # Add figure
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(fig_path), width=Inches(width_inches))
    
    # Space after figure
    add_paragraph(doc, "", font_size=6, spacing_after=4)


def parse_markdown_sections(content):
    """Parse markdown into a list of (type, level, text) blocks."""
    lines = content.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Heading
        if line.startswith("# "):
            blocks.append(("heading", 0, line[2:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("heading", 1, line[3:].strip()))
            i += 1
        elif line.startswith("### "):
            blocks.append(("heading", 2, line[4:].strip()))
            i += 1
        elif line.startswith("#### "):
            blocks.append(("heading", 3, line[5:].strip()))
            i += 1
        
        # Horizontal rule or table marker
        elif line.strip() == "---" or line.strip().startswith("|---"):
            i += 1
        
        # Table row
        elif line.strip().startswith("|") and line.strip().endswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            blocks.append(("table", 0, rows))
        
        # Empty line
        elif not line.strip():
            i += 1
        
        # Image reference
        elif line.strip().startswith("!["):
            blocks.append(("image", 0, line.strip()))
            i += 1
        
        # Regular paragraph
        else:
            para_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not lines[i].strip().startswith("|") and not lines[i].strip().startswith("---") and not lines[i].strip().startswith("!["):
                para_lines.append(lines[i])
                i += 1
            if para_lines:
                blocks.append(("paragraph", 0, "\n".join(para_lines)))
    
    return blocks


def clean_text(text):
    """Remove markdown formatting from text."""
    # Remove bold markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Replace LaTeX with simplified text
    text = re.sub(r'\$([^$]+?)\$', r'\1', text)
    # Remove backtick code
    text = re.sub(r'`([^`]+?)`', r'\1', text)
    return text


def parse_table_rows(rows):
    """Parse markdown table rows into structured data. Skips separator rows like |---|---|."""
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        # Skip separator rows
        if all(re.match(r'^[-: ]+$', c) for c in cells):
            continue
        parsed.append(cells)
    return parsed


def build_docx():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()
    
    doc = Document()
    
    # ─── Page setup ───
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # ─── Default style ───
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    
    blocks = parse_markdown_sections(content)
    
    figure_count = 0
    table_count = 0
    current_section = 0  # 0=front, 1=body, 2=appendix
    
    for block_type, level, text in blocks:
        # ─── Headings ───
        if block_type == "heading":
            if level == 0:
                # Title
                add_paragraph(doc, text, font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
            elif level == 1:
                if "附表" in text or "附录" in text:
                    current_section = 2
                add_heading_custom(doc, text, level=1)
            elif level == 2:
                if "摘要" in text:
                    current_section = 0
                elif "引言" in text:
                    current_section = 1
                add_heading_custom(doc, text, level=2)
            elif level == 3:
                add_heading_custom(doc, text, level=3)
            else:
                add_heading_custom(doc, text, level=4)
        
        # ─── Paragraph ───
        elif block_type == "paragraph":
            cleaned = clean_text(text)
            
            # Find figure references in the text
            fig_refs = []
            for fig_name in sorted(FIGURES.keys(), key=lambda x: -len(x)):
                if fig_name in cleaned:
                    fig_refs.append((cleaned.index(fig_name), fig_name, FIGURES[fig_name]))
            
            if fig_refs:
                # Sort by position
                fig_refs.sort(key=lambda x: x[0])
                
                # Add the paragraph text (without breaking)
                add_paragraph(doc, cleaned)
                
                # Add all referenced figures after the paragraph
                for _, fig_name, fig_path in fig_refs:
                    if fig_path.exists():
                        figure_count += 1
                        add_figure(doc, fig_path, fig_name, width_inches=5.0)
                    else:
                        add_paragraph(doc, f"[Missing: {fig_path.name}]", font_size=9, 
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_paragraph(doc, cleaned)
        
        # ─── Table ───
        elif block_type == "table":
            table_data = parse_table_rows(text)
            if len(table_data) < 2:
                continue
            
            table_count += 1
            n_rows = len(table_data)
            n_cols = max(len(r) for r in table_data)
            
            # Add table caption
            add_paragraph(doc, f"表 {table_count}.", font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
            
            # Create table
            tbl = doc.add_table(rows=n_rows, cols=n_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    if j < n_cols:
                        cell = tbl.cell(i, j)
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        run.font.size = Pt(8)
                        run.font.name = "宋体"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        
                        if i == 0:
                            run.bold = True
                        
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Reduce cell margins
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcMar = parse_xml(
                            f'<w:tcMar {nsdecls("w")}>'
                            f'<w:top w:w="30" w:type="dxa"/>'
                            f'<w:bottom w:w="30" w:type="dxa"/>'
                            f'<w:left w:w="60" w:type="dxa"/>'
                            f'<w:right w:w="60" w:type="dxa"/>'
                            f'</w:tcMar>'
                        )
                        tcPr.append(tcMar)
            
            # Apply three-line table style
            format_three_line_table(tbl)
            
            # Space after table
            add_paragraph(doc, "", font_size=6, spacing_after=4)
        
        # ─── Image (standalone) ───
        elif block_type == "image":
            pass
    
    # ─── Insert remaining figures into appropriate appendices ───
    # We'll add a figure appendix section with all unused figures
    used_figs = set()
    # Scan content for figure references
    for fig_name in FIGURES:
        if fig_name in content:
            used_figs.add(fig_name)
    
    unused_figs = {k: v for k, v in FIGURES.items() if k not in used_figs and v.exists()}
    if unused_figs:
        doc.add_page_break()
        add_heading_custom(doc, "附图", level=1)
        for i, (fig_name, fig_path) in enumerate(unused_figs.items(), 1):
            add_figure(doc, fig_path, f"附图 {i}. {fig_name}", width_inches=5.0)
    
    # Save
    doc.save(str(OUTPUT))
    print(f"DOCX saved: {OUTPUT}")
    print(f"  Figures embedded: {figure_count}")
    print(f"  Tables formatted: {table_count}")


if __name__ == "__main__":
    build_docx()
