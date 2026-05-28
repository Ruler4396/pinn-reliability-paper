"""
Apply text fixes from .md to user's formatted .docx, preserving formatting.
"""
from docx import Document
from copy import deepcopy
import re

def replace_in_docx(doc, old, new):
    """Replace text in all runs, preserving formatting. Handles text split across runs."""
    for p in doc.paragraphs:
        if old not in p.text:
            continue
        # Simple approach: if text fits in one run, replace there
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return True
        # Text split across runs: collect all text, replace, redistribute
        full = "".join(r.text for r in p.runs)
        if old in full:
            full = full.replace(old, new)
            # Put back into first run, clear others
            for i, r in enumerate(p.runs):
                if i == 0:
                    r.text = full
                else:
                    r.text = ""
            return True
    return False

def replace_in_tables(doc, old, new):
    """Replace in table cells."""
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)

doc = Document("paper_manuscript.docx")
count = 0

# ── English terms → Chinese ──
fixes = [
    # PDE names in body text
    ("Poisson's Equation", "泊松方程"),
    ("Navier-斯托克斯 方程", "纳维-斯托克斯方程"),
    ("Navier-斯托克斯", "纳维-斯托克斯"),
    # Remaining English
    ("case-specific", "针对特定案例的"),
    ("top-quartile", "上四分位"),
    ("Top-1/3", "前三分之一"),
    ("cluster bootstrap", "簇自助法"),
    ("RAR", "残差自适应重采样"),
    ("Dropout-based", "基于随机丢弃的"),
    ("Bayesian PINN", "贝叶斯 PINN"),
    # Training terminology
    ("curriculum learning", "课程学习"),
    # Statistics
    (" Wilson 置信区间", " 威尔逊置信区间"),
    (" Wilson 区间", " 威尔逊区间"),
    (" Spearman ", " 斯皮尔曼 "),
    ("Jaccard", "雅卡尔"),
    (" Cohen's $d_z$", " 科恩 $d_z$"),
    ("Cohen's $d_z$", "科恩 $d_z$"),
    # Mixed phrases
    ("Bootstrap", "自助法"),
    ("baseline", "基线配置"),
    ("stronger baselines", "更强的基线配置"),
    ("Adam 优化器", "Adam 优化器"),
    ("tanh 激活函数", "tanh 激活函数"),
    # Seeds  
    ("5 seeds", "5 个随机种子"),
    ("30 seeds ", "30 个种子 "),
    ("40 seeds", "40 个种子"),
    ("5-seed", "5 种子"),
    ("30-seed", "30 种子"),
    # Loss weights
    ("data=10, physics=1, boundary=10", "数据项权重10、物理项权重1、边界项权重10"),
    ("data=5, physics=2, boundary=15", "数据项权重5、物理项权重2、边界项权重15"),
    # Garbled patterns
    ("配对 自助法 加 科恩", "配对自助法配合科恩"),
    ("配对自助法 加 科恩", "配对自助法配合科恩"),
    ('"参考解-free"', "不依赖参考解"),
    ("参考解-free", "不依赖参考解"),
    ("优化器配置（, lr=", "Adam 优化器配置（学习率 "),
    ("， 激活函数", "，tanh 激活函数"),
    ("优化器，初始学习率", "Adam 优化器，初始学习率"),
    ("R-only", "仅综合分识别"),
    ("L2-only", "仅相对误差识别"),
    # Fix the broken appendix sentence
    ("附录的 展示了四个案例", "原始指标的跨系统对比图（见补充材料）"),
    ("（图 5）", "（图6-2a）"),
    ("（图 6）", "（图6-2b）"),
    # Old figure numbers
    ("（图5）", "（图6-2a）"),
    ("（图6）", "（图6-2b）"),
    ("7 counts → 4 training", "7 个格点变为训练稳定性主导 4 个"),
    ("physics_consistency", "物理约束"),
    ("training_stability", "训练稳定性"),
    ("numerical_accuracy", "数值精度"),
    ("structural_stability", "结构保真度"),
    # Double parentheses
    ("（（图6-2a，", "（图6-2a，"),
    ("（（图6-2b，", "（图6-2b，"),
    ("识别））", "识别）"),
    # PCA
    (" PCA ", " 主成分分析 "),
    ("PCA ", "主成分分析 "),
]

for old, new in fixes:
    if replace_in_docx(doc, old, new):
        count += 1
    replace_in_tables(doc, old, new)

# ── Fix table headers ──
# Table 1: Remove "Bayesian PINN" and "Dropout-based"  
# Table 2: Remove duplicate header row

# ── Remove orphan figure reference paragraph ──
for p in doc.paragraphs:
    if "图 3-(a)(b)(c) 分别展示" in p.text:
        p.clear()
        count += 1
        break

# ── Insert new figures ──
from pathlib import Path
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIG_DIR = Path("minimal_pinn") / "results" / "paper_figures" / "v4"

def insert_after(para_elem, new_elem):
    parent = para_elem.getparent()
    idx = list(parent).index(para_elem)
    parent.insert(idx + 1, new_elem)

# Find paragraphs where to insert figures
fig_to_insert = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if "**泊松方程。**" in t or "泊松方程。" in t:
        fig_to_insert.append((i, FIG_DIR / "fig5-1a.png", "图5-1a"))
    elif "**斯托克斯-泊肃叶流。**" in t:
        fig_to_insert.append((i, FIG_DIR / "fig5-1b.png", "图5-1b"))
    elif "**费希尔方程。**" in t:
        fig_to_insert.append((i, FIG_DIR / "fig5-1c.png", "图5-1c"))
    elif "**伯格斯方程。**" in t:
        fig_to_insert.append((i, FIG_DIR / "fig5-1d.png", "图5-1d"))

# Insert figures (reverse order)
added_figs = 0
for pi, fp, cap in reversed(fig_to_insert):
    if not fp.exists():
        continue
    ref_p = doc.paragraphs[pi]
    
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run(cap).font.size = Pt(9)
    ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(str(fp), width=Inches(6.0))
    sp = doc.add_paragraph()
    
    # Move elements
    for e in [cp._element, ip._element, sp._element]:
        e.getparent().remove(e)
        insert_after(ref_p._element, e)
    added_figs += 1

print(f"Text fixes: {count}")
print(f"Figures inserted: {added_figs}")

doc.save("paper_manuscript.docx")
print("Saved: paper_manuscript.docx")
