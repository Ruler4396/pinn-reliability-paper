"""Final postprocess: insert all figures into pandoc docx by paragraph position."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT = Path(__file__).resolve().parent.parent
FIG_DIR = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v4"
MORPH_DIR = PROJECT / "minimal_pinn" / "results" / "analysis" / "divergence_morphology_v1"
FIG_DIR_V2 = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v2"
INPUT = PROJECT / "paper_manuscript_pandoc.docx"
OUTPUT = PROJECT / "paper_manuscript_final.docx"

def insert_after(elem, new_elem):
    parent = elem.getparent()
    parent.insert(list(parent).index(elem) + 1, new_elem)

def add_figure_after(doc, paragraph_idx, fig_path, caption=""):
    """Add figure (caption + image) after paragraph at given index."""
    ref = doc.paragraphs[paragraph_idx]
    
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption).font.size = Pt(9)
    img = doc.add_paragraph(); img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img.add_run().add_picture(str(fig_path), width=Inches(6.0))
    sp = doc.add_paragraph()
    
    ce, ie, se = cap._element, img._element, sp._element
    for e in [ce, ie, se]: e.getparent().remove(e)
    for e in [se, ie, ce]: insert_after(ref._element, e)
    return True

def main():
    doc = Document(str(INPUT))
    
    # Find key paragraphs by content
    para_map = {}  # keyword → paragraph index
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        for kw in ["5.1 四案例的退化全貌", "**泊松方程**", "**斯托克斯-泊肃叶流**",
                    "**费希尔方程**", "**伯格斯方程**", "5.2 三系统的定量梯度",
                    "6.1 主导维度分布", "6.2 消融分析", "6.3 差集工况的形态对比",
                    "图6-2a", "图6-2b", "图A-1", "图A-2", "图A-3",
                    "7.1 校准敏感性与反循环检查", "7.2 阈值可移植性"]:
            if kw in t and kw not in para_map:
                para_map[kw] = i
    
    print(f"Found {len(para_map)} anchor paragraphs")
    
    # Insert figures at specific locations
    inserts = []
    
    # Section 5.1 - individual case phase maps after each case
    for case_kw, fig_names in [
        ("**泊松方程**", ["fig5-1a.png", "fig5-2a.png"]),
        ("**斯托克斯-泊肃叶流**", ["fig5-1b.png", "fig5-2b.png"]),
        ("**费希尔方程**", ["fig5-1c.png", "fig5-2c.png"]),
        ("**伯格斯方程**", ["fig5-1d.png", "fig5-2d.png"]),
    ]:
        if case_kw in para_map:
            for fn in fig_names:
                fp = FIG_DIR / fn
                if fp.exists():
                    inserts.append((para_map[case_kw], fp, fn.replace(".png","").replace("fig","图")))

    # Section 5.2 - three-system probability heatmaps
    if "5.2 三系统的定量梯度" in para_map:
        for fn in ["fig5-3a.png","fig5-3b.png","fig5-3c.png","fig5-4.png"]:
            fp = FIG_DIR / fn
            if fp.exists():
                inserts.append((para_map["5.2 三系统的定量梯度"], fp, fn.replace(".png","").replace("fig","图")))

    # Section 6.1 - dominant dimension (no prepopulated fig)
    # Section 6.2 - ablation
    if "6.2 消融分析" in para_map:
        fp = FIG_DIR / "fig6-1.png"
        if fp.exists():
            inserts.append((para_map["6.2 消融分析"], fp, "图6-1"))

    # Section 6.3 - divergence morphology  
    if "图6-2a" in para_map:
        fp = MORPH_DIR / "burgers_R-worst_not_L2-worst.png"
        if fp.exists():
            inserts.append((para_map["图6-2a"], fp, "图6-2a"))
    if "图6-2b" in para_map:
        fp = MORPH_DIR / "burgers_L2-worst_not_R-worst.png"
        if fp.exists():
            inserts.append((para_map["图6-2b"], fp, "图6-2b"))

    # Section 7 - validation
    for kw, fn in [
        ("7.1 校准敏感性与反循环检查", "fig7-1.png"),
        ("7.1 校准敏感性与反循环检查", "fig7-2.png"),
        ("7.2 阈值可移植性", "fig7-3.png"),
    ]:
        if kw in para_map:
            fp = FIG_DIR / fn
            if fp.exists():
                inserts.append((para_map[kw], fp, fn.replace(".png","").replace("fig","图")))

    # Appendix
    for kw in ["图A-1","图A-2","图A-3"]:
        if kw in para_map:
            labels = {"图A-1":"a","图A-2":"b","图A-3":"c"}
            fp = FIG_DIR_V2 / f"fig_{labels[kw]}{kw.replace('图A-','')}_calibration_sensitivity.png" if "A-1" in kw else \
                 FIG_DIR_V2 / f"fig_{labels[kw]}{kw.replace('图A-','')}_anti_circularity.png" if "A-2" in kw else \
                 FIG_DIR_V2 / f"fig_{labels[kw]}{kw.replace('图A-','')}_baseline_failure.png"
            if fp.exists():
                inserts.append((para_map[kw], fp, kw))

    # Apply inserts, reversed
    inserts.sort(key=lambda x: -x[0])
    done = 0
    for pi, fp, cap in inserts:
        try:
            add_figure_after(doc, pi, fp, cap)
            done += 1
        except Exception as e:
            print(f"  Failed {cap}: {e}")
    
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}, {done} figures inserted")

if __name__ == "__main__":
    main()
