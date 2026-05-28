"""
Build docx with proper OMML equations instead of plain-text formulas.
Handles: inline $...$, display $$...$$, fractions, subscripts, sums, norms, Greek letters.
"""
from __future__ import annotations

import re, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

PROJECT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT / "paper_manuscript_zh.md"
FIG_DIR = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v2"
MORPH_DIR = PROJECT / "minimal_pinn" / "results" / "analysis" / "divergence_morphology_v1"
OUTPUT = PROJECT / "paper_manuscript_zh.docx"

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def make_math_elem(tag):
    """Create an element in the math namespace."""
    from lxml import etree
    return etree.SubElement(etree.Element("dummy"), f"{{{MATH_NS}}}{tag}")

def make_math_elem_root(tag):
    """Create a root element in the math namespace."""
    from lxml import etree
    return etree.Element(f"{{{MATH_NS}}}{tag}", nsmap={"m": MATH_NS})

def make_omath():
    return OxmlElement(Mtag("oMath"))

def make_run(text="", italic=False):
    r = OxmlElement(Mtag("r"))
    if italic:
        rPr = OxmlElement(Mtag("rPr"))
        sty = OxmlElement(Mtag("sty"))
        sty.set(Mtag("val"), "i")
        rPr.append(sty)
        r.append(rPr)
    t = OxmlElement(Mtag("t"))
    t.text = text
    t.set("xml:space", "preserve")
    r.append(t)
    return r

def make_frac(num_elems, den_elems):
    f = OxmlElement(Mtag("f"))
    num = OxmlElement(Mtag("num"))
    for e in num_elems: num.append(e)
    den = OxmlElement(Mtag("den"))
    for e in den_elems: den.append(e)
    f.append(num); f.append(den)
    return f

def make_sub(elems, sub_text):
    s = OxmlElement(Mtag("sSub"))
    e = OxmlElement(Mtag("e"))
    for el in elems: e.append(el)
    sub = OxmlElement(Mtag("sub"))
    sub.append(make_run(str(sub_text)))
    s.append(e); s.append(sub)
    return s

def make_sup(elems, sup_text):
    s = OxmlElement(Mtag("sSup"))
    e = OxmlElement(Mtag("e"))
    for el in elems: e.append(el)
    sup = OxmlElement(Mtag("sup"))
    sup.append(make_run(str(sup_text)))
    s.append(e); s.append(sup)
    return s

def make_rad(elems):
    r = OxmlElement(Mtag("rad"))
    deg = OxmlElement(Mtag("deg"))
    deg.append(make_run("2"))
    e = OxmlElement(Mtag("e"))
    for el in elems: e.append(el)
    r.append(deg); r.append(e)
    return r

def make_norm(elems, sub=""):
    """Create ||x||_2 style norm."""
    d = OxmlElement(Mtag("d"))
    dPr = OxmlElement(Mtag("dPr"))
    beg = OxmlElement(Mtag("begChr"))
    beg.set(Mtag("val"), "‖")
    dPr.append(beg)
    sep = OxmlElement(Mtag("sepChr"))
    sep.set(Mtag("val"), "‖")
    dPr.append(sep)
    end = OxmlElement(Mtag("endChr"))
    end.set(Mtag("val"), "‖")
    dPr.append(end)
    d.append(dPr)
    e = OxmlElement(Mtag("e"))
    for el in elems: e.append(el)
    d.append(e)
    if sub:
        return make_sub([d], sub)
    return d

def make_group(elems):
    """Group with parentheses."""
    d = OxmlElement(Mtag("d"))
    dPr = OxmlElement(Mtag("dPr"))
    beg = OxmlElement(Mtag("begChr")); beg.set(Mtag("val"), "("); dPr.append(beg)
    sep = OxmlElement(Mtag("sepChr")); sep.set(Mtag("val"), "|"); dPr.append(sep)
    end = OxmlElement(Mtag("endChr")); end.set(Mtag("val"), ")"); dPr.append(end)
    d.append(dPr)
    e = OxmlElement(Mtag("e"))
    for el in elems: e.append(el)
    d.append(e)
    return d

def make_sum(lo_text, hi_text, body_elems):
    """Summation: Σ_{lo}^{hi} body."""
    nary = OxmlElement(Mtag("nary"))
    nPr = OxmlElement(Mtag("naryPr"))
    chr_el = OxmlElement(Mtag("chr")); chr_el.set(Mtag("val"), "∑"); nPr.append(chr_el)
    lim = OxmlElement(Mtag("limLoc")); lim.set(Mtag("val"), "undOvr"); nPr.append(lim)
    nary.append(nPr)
    sub = OxmlElement(Mtag("sub")); sub.append(make_run(str(lo_text))); nary.append(sub)
    sup = OxmlElement(Mtag("sup")); sup.append(make_run(str(hi_text))); nary.append(sup)
    e = OxmlElement(Mtag("e"))
    for el in body_elems: e.append(el)
    nary.append(e)
    return nary

def make_barrier(elems):
    """Absolute value |x|."""
    d = OxmlElement(Mtag("d"))
    dPr = OxmlElement(Mtag("dPr"))
    beg = OxmlElement(Mtag("begChr")); beg.set(Mtag("val"), "|"); dPr.append(beg)
    sep = OxmlElement(Mtag("sepChr")); sep.set(Mtag("val"), "|"); dPr.append(sep)
    end = OxmlElement(Mtag("endChr")); end.set(Mtag("val"), "|"); dPr.append(end)
    d.append(dPr)
    e = OxmlElement(Mtag("e"))
    for el in elems: e.append(el)
    d.append(e)
    return d

def make_hat(elems):
    acc = OxmlElement(Mtag("acc"))
    accPr = OxmlElement(Mtag("accPr"))
    chr_el = OxmlElement(Mtag("chr")); chr_el.set(Mtag("val"), "̂"); accPr.append(chr_el)
    acc.append(accPr)
    e = OxmlElement(Mtag("e"))
    for el in elems: e.append(el)
    acc.append(e)
    return acc

def parse_inline_math(expr):
    """Convert inline LaTeX to OMML elements. Returns list of OMML elements."""
    elems = []
    i = 0
    while i < len(expr):
        # Subscript: x_{text}
        if i + 1 < len(expr) and expr[i] == '{':
            # Find matching }
            depth = 1; j = i + 1
            while j < len(expr) and depth > 0:
                if expr[j] == '{': depth += 1
                elif expr[j] == '}': depth -= 1
                j += 1
            inner = expr[i+1:j-1]
            # Check if preceding char is _ or ^
            if elems and i > 0 and expr[i-1] == '_':
                prev = elems.pop()
                parsed_inner = parse_inline_math(inner)
                elems.append(make_sub(parsed_inner if parsed_inner else [make_run(inner)], ""))
            elif elems and i > 0 and expr[i-1] == '^':
                prev = elems.pop()
                parsed_inner = parse_inline_math(inner)
                elems.append(make_sup([make_run(prev.text)] if hasattr(prev,'text') else [prev], inner))
            else:
                parsed = parse_inline_math(inner)
                elems.extend(parsed)
            i = j
            continue
        
        # Commands: \mathrm{text}, \mathcal{X}, \sigma, etc.
        if expr[i] == '\\':
            j = i + 1
            while j < len(expr) and expr[j].isalpha():
                j += 1
            cmd = expr[i+1:j]
            
            if cmd == 'sigma': elems.append(make_run("σ"))
            elif cmd == 'epsilon': elems.append(make_run("ε"))
            elif cmd == 'mathcal':
                # Next char should be {L} or similar
                if j < len(expr) and expr[j] == '{':
                    k = expr.index('}', j)
                    letter = expr[j+1:k]
                    script_map = {'L': 'ℒ', 'F': 'ℱ', 'B': 'ℬ', 'D': '𝒟', 'T': '𝒯', 'S': '𝒮'}
                    elems.append(make_run(script_map.get(letter, letter), italic=True))
                    i = k + 1
                    continue
            elif cmd == 'mathrm':
                if j < len(expr) and expr[j] == '{':
                    k = expr.index('}', j)
                    text = expr[j+1:k]
                    # Map common subscripts
                    text = text.replace('obs', 'obs').replace('col', 'col')
                    for rch in text:
                        elems.append(make_run(rch))
                    i = k + 1
                    continue
            elif cmd == 'mathbf': pass
            elif cmd == 'hat':
                if j < len(expr) and expr[j] == '{':
                    k = expr.index('}', j)
                    inner = expr[j+1:k]
                    inner_elems = parse_inline_math(inner)
                    elems.append(make_hat(inner_elems if inner_elems else [make_run(inner)]))
                    i = k + 1
                    continue
            elif cmd == 'arg':
                # \arg\min
                pass
            elif cmd == 'min': elems.append(make_run("min"))
            elif cmd == 'max': elems.append(make_run("max"))
            elif cmd == 'exp': elems.append(make_run("exp"))
            elif cmd == 'ln': elems.append(make_run("ln"))
            elif cmd == 'log': elems.append(make_run("log"))
            elif cmd == 'sin': elems.append(make_run("sin"))
            elif cmd == 'pi': elems.append(make_run("π"))
            elif cmd == 'nu': elems.append(make_run("ν"))
            elif cmd == 'mu': elems.append(make_run("μ"))
            elif cmd == 'theta': elems.append(make_run("θ"))
            elif cmd == 'infty': elems.append(make_run("∞"))
            elif cmd == 'cdot': elems.append(make_run("·"))
            elif cmd == 'times': elems.append(make_run("×"))
            elif cmd == 'pm': elems.append(make_run("±"))
            elif cmd == 'le': elems.append(make_run("≤"))
            elif cmd == 'ge': elems.append(make_run("≥"))
            elif cmd == 'to': elems.append(make_run("→"))
            elif cmd == 'sim': elems.append(make_run("∼"))
            elif cmd == 'in': elems.append(make_run("∈"))
            elif cmd == 'partial': elems.append(make_run("∂"))
            elif cmd == 'Omega': elems.append(make_run("Ω"))
            elif cmd == 'ell': elems.append(make_run("ℓ"))
            elif cmd == 'sum': elems.append(make_run("∑"))
            elif cmd == 'prod': elems.append(make_run("∏"))
            elif cmd == 'int': elems.append(make_run("∫"))
            elif cmd == 'left': pass
            elif cmd == 'right': pass
            elif cmd == 'langle': elems.append(make_run("⟨"))
            elif cmd == 'rangle': elems.append(make_run("⟩"))
            elif cmd == 'ldots': elems.append(make_run("…"))
            elif cmd == 'quad': elems.append(make_run("  "))
            elif cmd == 'bar':
                if j < len(expr) and expr[j] == '{':
                    k = expr.index('}', j)
                    inner = expr[j+1:k]
                    elems.append(make_run(f"{inner}̄"))
                    i = k + 1
                    continue
            elif cmd == 'dot':
                if j < len(expr) and expr[j] == '{':
                    k = expr.index('}', j)
                    inner = expr[j+1:k]
                    elems.append(make_run(f"{inner}̇"))
                    i = k + 1
                    continue
            elif cmd == 'mathrm': pass  # handled above
            else:
                elems.append(make_run(cmd))
            i = j
            continue
        
        # Underscore subscript
        if expr[i] == '_':
            if i + 1 < len(expr):
                j = i + 1
                if expr[j] == '{':
                    k = expr.index('}', j)
                    sub_text = expr[j+1:k]
                else:
                    k = j + 1
                    sub_text = expr[j:k]
                if elems:
                    prev = elems.pop()
                    inner = parse_inline_math(sub_text)
                    elems.append(make_sub([make_run(prev.text)] if hasattr(prev,'text') else [prev], sub_text))
                i = k
                continue
        
        # Caret superscript
        if expr[i] == '^':
            if i + 1 < len(expr):
                j = i + 1
                if expr[j] == '{':
                    k = expr.index('}', j)
                    sup_text = expr[j+1:k]
                elif expr[j] == '\\':
                    # \star, etc
                    k = j + 2
                    while k < len(expr) and expr[k].isalpha(): k += 1
                    sup_text = expr[j:k]
                else:
                    k = j + 1
                    sup_text = expr[j:k]
                if elems:
                    prev = elems.pop()
                    elems.append(make_sup([make_run(prev.text)] if hasattr(prev,'text') else [prev], sup_text))
                i = k
                continue
        
        # Pipe character - treat as norm barrier
        if expr[i] == '|':
            elems.append(make_run("|"))
            i += 1
            continue
        
        # Regular character
        elems.append(make_run(expr[i]))
        i += 1
    
    return elems


def formula_to_omath(formula):
    """Convert a LaTeX formula string to an oMath OMML element."""
    om = make_omath()
    
    # Remove $$ markers for display
    formula = formula.replace("$$", "").strip()
    
    # Try to handle fractions: \frac{num}{den}
    frac_match = re.search(r'\\frac\{(.+?)\}\{(.+?)\}', formula)
    if frac_match:
        prefix = formula[:frac_match.start()]
        suffix = formula[frac_match.end():]
        num = frac_match.group(1)
        den = frac_match.group(2)
        
        if prefix.strip():
            for e in parse_inline_math(prefix): om.append(e)
        
        num_elems = parse_inline_math(num)
        den_elems = parse_inline_math(den)
        om.append(make_frac(num_elems, den_elems))
        
        if suffix.strip():
            for e in parse_inline_math(suffix): om.append(e)
        return om
    
    # Try to handle sums: \sum_{lo}^{hi} body
    sum_match = re.search(r'\\sum_\{([^}]+)\}\^\{([^}]+)\}', formula)
    if sum_match:
        prefix = formula[:sum_match.start()]
        suffix = formula[sum_match.end():]
        lo = sum_match.group(1)
        hi = sum_match.group(2)
        body = suffix.strip()
        if prefix.strip():
            for e in parse_inline_math(prefix): om.append(e)
        om.append(make_sum(lo, hi, parse_inline_math(body)))
        return om
    
    # Default: parse as inline
    for e in parse_inline_math(formula):
        om.append(e)
    return om


def insert_omath(paragraph, om):
    """Insert oMath element into a paragraph."""
    # Find or create a run to insert after
    last_run = None
    for r in paragraph.runs:
        last_run = r
    if last_run is None:
        last_run = paragraph.add_run("")
    last_run._element.addnext(om)


def add_paragraph_formatted(doc, text, style="Normal", bold=False, font_size=10.5, 
                            alignment=None, spacing_after=6):
    """Add paragraph with proper formatting and formula handling."""
    # Replace inline $...$ with formula markers
    # We'll split text by inline formulas and add runs/equations
    parts = re.split(r'(\$[^$]+\$)', text)
    
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    pf.line_spacing = 1.5
    if alignment:
        p.alignment = alignment
    
    for part in parts:
        if part.startswith("$") and part.endswith("$"):
            formula = part[1:-1].strip()
            if formula:
                om = formula_to_omath(formula)
                insert_omath(p, om)
        else:
            if part.strip():
                run = p.add_run(part)
                run.font.size = Pt(font_size)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if bold:
                    run.bold = True
    return p


def add_heading_custom(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 0:
        run.font.size = Pt(16)
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(10.5)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True
    return h


def add_display_formula(doc, formula):
    """Add a display formula ($$...$$) as a centered paragraph."""
    formula = formula.replace("$$", "").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    om = formula_to_omath(formula)
    insert_omath(p, om)
    return p


def add_figure(doc, fig_path, caption):
    if not fig_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"[Figure not found: {fig_path.name}]").font.size = Pt(9)
        return
    # Caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption)
    cr.font.size = Pt(9)
    cr.font.name = "宋体"
    cr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # Image
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = ip.add_run()
    ir.add_picture(str(fig_path), width=Inches(6.0))
    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def format_three_line_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    # Clear all borders
    for cell in table._cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        for bn in ["top","left","bottom","right","insideH","insideV"]:
            border = parse_xml(f'<w:{bn} {nsdecls("w")} w:val="nil"/>')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    # Top and bottom borders on table
    tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:left w:val="nil"/><w:right w:val="nil"/>'
        f'<w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        f'</w:tblBorders>')
    tblPr.append(tblBorders)
    # Bottom border on header row
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)
        bottom = parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="6" w:space="0" w:color="000000"/>')
        tcBorders.append(bottom)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def build_docx():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()
    
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    
    # Parse content line by line
    lines = content.split("\n")
    i = 0
    figure_count = 0
    table_count = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Empty line
        if not stripped:
            i += 1
            continue
        
        # Display formula: $$...$$
        if stripped.startswith("$$"):
            formula = ""
            i += 1
            while i < len(lines) and not lines[i].strip().endswith("$$"):
                formula += lines[i] + "\n"
                i += 1
            if i < len(lines):
                formula += lines[i]
                i += 1
            add_display_formula(doc, formula)
            continue
        
        # Heading
        if stripped.startswith("# "):
            add_heading_custom(doc, stripped[2:], 0)
            i += 1
        elif stripped.startswith("## "):
            add_heading_custom(doc, stripped[3:], 1)
            i += 1
        elif stripped.startswith("### "):
            add_heading_custom(doc, stripped[4:], 2)
            i += 1
        elif stripped.startswith("#### "):
            add_heading_custom(doc, stripped[5:], 3)
            i += 1
        
        # Horizontal rule → skip
        elif stripped == "---":
            i += 1
        
        # Table
        elif stripped.startswith("|") and stripped.endswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            # Parse and add table
            parsed = []
            for row in rows:
                cells = [c.strip() for c in row.strip("|").split("|")]
                if all(re.match(r'^[-: ]+$', c) for c in cells):
                    continue
                parsed.append(cells)
            if len(parsed) >= 2:
                table_count += 1
                n_rows = len(parsed)
                n_cols = max(len(r) for r in parsed)
                tbl = doc.add_table(rows=n_rows, cols=n_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row_data in enumerate(parsed):
                    for ci, cell_text in enumerate(row_data):
                        if ci < n_cols:
                            cell = tbl.cell(ri, ci)
                            p = cell.paragraphs[0]
                            r = p.add_run(cell_text)
                            r.font.size = Pt(8)
                            r.font.name = "宋体"
                            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                            if ri == 0:
                                r.bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                format_three_line_table(tbl)
                # Spacer
                sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
        
        # Regular paragraph
        else:
            para_text = stripped
            i += 1
            # Collect continuation lines
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not lines[i].strip().startswith("$$") and not lines[i].strip().startswith("|") and lines[i].strip() != "---":
                para_text += " " + lines[i].strip()
                i += 1
            
            # Check for figure references and insert after paragraph
            fig_refs = []
            for fn in sorted(FIGURES.keys(), key=lambda x: -len(x)):
                if fn in para_text and FIGURES[fn].exists():
                    fig_refs.append(fn)
            
            add_paragraph_formatted(doc, para_text)
            
            for fn in fig_refs:
                figure_count += 1
                add_figure(doc, FIGURES[fn], fn)
    
    doc.save(str(OUTPUT))
    print(f"DOCX saved: {OUTPUT}")
    print(f"  Figures: {figure_count}, Tables: {table_count}")


if __name__ == "__main__":
    build_docx()
