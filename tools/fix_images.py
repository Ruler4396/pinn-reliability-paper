"""
Properly insert v4 figures into docx, preserving relationships correctly.
The issue: add_picture() in temp docs doesn't transfer relationships.
Fix: add pictures directly in the main doc.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

DOC = Document("paper_manuscript.docx")
FIG_DIR = Path("minimal_pinn/results/paper_figures/v4")
MORPH_DIR = Path("minimal_pinn/results/analysis/divergence_morphology_v1")
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
body = DOC.element.body

# ── Step 1: Remove ALL existing images and captions ──
removed = 0
for elem in list(body):
    # Check for image (drawing)
    has_drawing = elem.find(f'.//{{{NS}}}drawing') is not None
    # Check for figure caption
    texts = elem.findall(f'.//{{{NS}}}t')
    full_text = "".join(t.text or "" for t in texts)
    is_caption = "图5-" in full_text or "图6-" in full_text or "图7-" in full_text
    
    if has_drawing or (is_caption and len(full_text) < 100):
        body.remove(elem)
        removed += 1

print(f"Removed {removed} old image/caption elements")

# ── Step 2: Insert new figures directly in main doc ──
insertions = {
    # (anchor_text_fragment, [(caption, filepath), ...])
    "泊松方程。" : [
        ("图5-1a 泊松方程：相对L2误差分布", FIG_DIR/"fig5-1a.png"),
    ],
    "斯托克斯-泊肃叶流。" : [
        ("图5-1b 斯托克斯-泊肃叶流：相对L2误差分布", FIG_DIR/"fig5-1b.png"),
    ],
    "费希尔方程。" : [
        ("图5-1c 费希尔方程：相对L2误差分布", FIG_DIR/"fig5-1c.png"),
    ],
    "伯格斯方程。" : [
        ("图5-1d 伯格斯方程：相对L2误差分布", FIG_DIR/"fig5-1d.png"),
    ],
    "表2. 三系统概率边界定量梯度。" : [
        ("图5-3a 斯托克斯-泊肃叶流：概率边界越界率", FIG_DIR/"fig5-3a.png"),
        ("图5-3b 费希尔方程：概率边界越界率", FIG_DIR/"fig5-3b.png"),
        ("图5-3c 伯格斯方程：概率边界越界率", FIG_DIR/"fig5-3c.png"),
        ("图5-4 边界关键点越界率的威尔逊置信区间（左：伯格斯，右：费希尔）", FIG_DIR/"fig5-4.png"),
    ],
    "6.2 消融分析" : [
        ("图6-1 消融对比：完整四维R与简化指标的跨种子排序一致性", FIG_DIR/"fig6-1.png"),
    ],
    "图6-2a" : [
        ("图6-2a 伯格斯方程：R标记为最差但rel2未标记的工况", MORPH_DIR/"burgers_R-worst_not_L2-worst.png"),
    ],
    "图6-2b" : [
        ("图6-2b 伯格斯方程：rel2标记为最差但R未标记的工况", MORPH_DIR/"burgers_L2-worst_not_R-worst.png"),
    ],
    "7.1 校准敏感性与反循环检查" : [
        ("图7-1 校准敏感性：不同分位点配置下的主导维度分布", FIG_DIR/"fig7-1.png"),
        ("图7-2 反循环校验：留出评估中主导维度的一致率", FIG_DIR/"fig7-2.png"),
    ],
    "7.2 阈值可移植性与跨变体一致性" : [
        ("图7-3 伯格斯方程：最安全点越界率随阈值倍数的变化", FIG_DIR/"fig7-3.png"),
    ],
}

def make_caption(text):
    """Create a centered caption paragraph element."""
    p = etree.Element(f"{{{NS}}}p")
    pPr = etree.SubElement(p, f"{{{NS}}}pPr")
    jc = etree.SubElement(pPr, f"{{{NS}}}jc")
    jc.set(f"{{{NS}}}val", "center")
    r = etree.SubElement(p, f"{{{NS}}}r")
    rPr = etree.SubElement(r, f"{{{NS}}}rPr")
    sz = etree.SubElement(rPr, f"{{{NS}}}sz")
    sz.set(f"{{{NS}}}val", "18")
    t = etree.SubElement(r, f"{{{NS}}}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p

def make_image_paragraph(fig_path):
    """Create an image paragraph by adding picture directly to main doc."""
    # Use python-docx to add the picture properly
    from docx import Document as D2
    # Create the paragraph element with the image
    p = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(6.0))
    # Get the element, remove from end of doc, return
    elem = p._element
    body.remove(elem)
    return elem

total = 0
for anchor_text, fig_list in insertions.items():
    # Find the anchor paragraph
    anchor_idx = None
    for i, p in enumerate(DOC.paragraphs):
        if p.text and anchor_text in p.text:
            # For case headers, prefer the ones in section 5.1 (P >= 100)
            if anchor_text.endswith("。") and i < 100:
                continue  # Skip Section 4.4 version, use Section 5.1 version
            anchor_idx = i
            break
    
    if anchor_idx is None:
        print(f"SKIP: anchor not found: {anchor_text[:40]}")
        continue
    
    anchor_elem = DOC.paragraphs[anchor_idx]._element
    anchor_pos = list(body).index(anchor_elem)
    
    print(f"P{anchor_idx}: {anchor_text[:40]} -> {len(fig_list)} figures")
    
    for caption, fig_path in fig_list:
        if not fig_path.exists():
            print(f"  MISSING: {fig_path}")
            continue
        
        cap_elem = make_caption(caption)
        img_elem = make_image_paragraph(fig_path)
        
        body.insert(anchor_pos + 1, img_elem)
        body.insert(anchor_pos + 1, cap_elem)
        total += 1

DOC.save("paper_manuscript.docx")
print(f"\nInserted {total} figures. Saved.")
