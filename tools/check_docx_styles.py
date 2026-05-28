import docx

d = docx.Document("paper_manuscript_zh.docx")

# Check font colors
print("=== Font colors in first 10 paragraphs ===")
for i, p in enumerate(d.paragraphs[:10]):
    for j, r in enumerate(p.runs):
        font = r.font
        color_str = str(font.color.rgb) if font.color and font.color.rgb else "none"
        print(f"P{i} R{j}: color={color_str}, bold={font.bold}, size={font.size}")

# Check table cell background (shading)
print("\n=== Table 1 cell backgrounds ===")
if d.tables:
    tbl = d.tables[0]
    for ri, row in enumerate(tbl.rows[:2]):
        for ci, cell in enumerate(row.cells[:2]):
            # Check shading
            tcPr = cell._tc.tcPr
            shading = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd") if tcPr is not None else None
            if shading is not None:
                fill = shading.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                color = shading.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color")
                print(f"  Cell({ri},{ci}): fill={fill}, color={color}")
            else:
                print(f"  Cell({ri},{ci}): no shading")

# Check paragraph-level shading
print("\n=== Paragraph shading ===")
for i, p in enumerate(d.paragraphs[:5]):
    pPr = p._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is not None:
        shd = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
        if shd is not None:
            fill = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
            print(f"P{i}: background fill={fill}")

print("\nDone. python-docx can read: font color, font size, bold, cell shading, paragraph shading.")
