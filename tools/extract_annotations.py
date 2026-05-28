"""Extract user's docx annotations: red text = AI flavor, yellow bg = awkward sentences."""
import docx

d = docx.Document("paper_manuscript_zh.docx")
NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Red-marked text (AI flavor)
print("=== RED-MARKED (AI flavor) ===")
for pi, p in enumerate(d.paragraphs):
    reds = []
    for r in p.runs:
        c = str(r.font.color.rgb) if r.font.color and r.font.color.rgb else ""
        if c == "EE0000" and r.text.strip():
            reds.append(r.text.strip())
    if reds:
        print(f"\nP{pi}:")
        for rt in reds:
            print(f"  [RED] {rt}")

# Yellow-highlighted (awkward sentences)
print("\n\n=== YELLOW-HIGHLIGHTED (awkward) ===")
for pi, p in enumerate(d.paragraphs):
    highlighted = False
    for r in p.runs:
        rPr = r._element.find(f"{NS}rPr")
        if rPr is not None:
            hl = rPr.find(f"{NS}highlight")
            if hl is not None:
                highlighted = True
                break
    if highlighted and p.text.strip():
        print(f"\nP{pi}:")
        print(f"  [BG] {p.text.strip()[:200]}")
