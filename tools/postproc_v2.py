"""Post-process pandoc .docx: insert images after figure references."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJECT = Path(__file__).resolve().parent.parent
FIG_DIR_V4 = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v4"
FIG_DIR_V2 = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v2"
MORPH_DIR = PROJECT / "minimal_pinn" / "results" / "analysis" / "divergence_morphology_v1"
INPUT = PROJECT / "paper_manuscript_pandoc.docx"
OUTPUT = PROJECT / "paper_manuscript_zh.docx"

FIG_MAP = {
    "图5-1a": [FIG_DIR_V4 / "fig5-1a.png"],
    "图5-1b": [FIG_DIR_V4 / "fig5-1b.png"],
    "图5-1c": [FIG_DIR_V4 / "fig5-1c.png"],
    "图5-1d": [FIG_DIR_V4 / "fig5-1d.png"],
    "图5-2a": [FIG_DIR_V4 / "fig5-2a.png"],
    "图5-2b": [FIG_DIR_V4 / "fig5-2b.png"],
    "图5-2c": [FIG_DIR_V4 / "fig5-2c.png"],
    "图5-2d": [FIG_DIR_V4 / "fig5-2d.png"],
    "图5-3a": [FIG_DIR_V4 / "fig5-3a.png"],
    "图5-3b": [FIG_DIR_V4 / "fig5-3b.png"],
    "图5-3c": [FIG_DIR_V4 / "fig5-3c.png"],
    "图5-4": [FIG_DIR_V4 / "fig5-4.png"],
    "图6-1": [FIG_DIR_V4 / "fig6-1.png"],
    "图6-2a": [MORPH_DIR / "burgers_R-worst_not_L2-worst.png"],
    "图6-2b": [MORPH_DIR / "burgers_L2-worst_not_R-worst.png"],
    "图7-1": [FIG_DIR_V4 / "fig7-1.png"],
    "图7-2": [FIG_DIR_V4 / "fig7-2.png"],
    "图7-3": [FIG_DIR_V4 / "fig7-3.png"],
    "图A-1": [FIG_DIR_V2 / "fig_a1_calibration_sensitivity.png"],
    "图A-2": [FIG_DIR_V2 / "fig_a2_anti_circularity.png"],
    "图A-3": [FIG_DIR_V2 / "fig_a3_baseline_failure.png"],
}

def insert_after(para_elem, new_elem):
    """Insert new_elem after para_elem in the XML tree."""
    parent = para_elem.getparent()
    idx = list(parent).index(para_elem)
    parent.insert(idx + 1, new_elem)

def main():
    doc = Document(str(INPUT))
    
    # Find figure reference paragraphs
    inserted_names = set()
    inserts = []
    for i, p in enumerate(doc.paragraphs):
        for name, paths in FIG_MAP.items():
            if name in p.text and name not in inserted_names:
                # Check all paths exist
                valid_paths = [fp for fp in paths if fp.exists()]
                if valid_paths:
                    inserts.append((i, name, valid_paths))
                    inserted_names.add(name)
    
    print(f"Found {len(inserts)} unique figure references")
    
    for para_idx, fig_name, fig_paths in reversed(inserts):
        ref_para = doc.paragraphs[para_idx]
        
        for fi, fig_path in enumerate(fig_paths):
            label = fig_name if len(fig_paths) == 1 else f"{fig_name}-({chr(ord('a')+fi)})"
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(label)
            cap_run.font.size = Pt(9)
            
            img = doc.add_paragraph()
            img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img.add_run().add_picture(str(fig_path), width=Inches(6.0))
            
            sp = doc.add_paragraph()
            
            cp_elem = cap._element; cap._element.getparent().remove(cp_elem)
            im_elem = img._element; img._element.getparent().remove(im_elem)
            sp_elem = sp._element; sp._element.getparent().remove(sp_elem)
            
            insert_after(ref_para._element, cp_elem)
            insert_after(ref_para._element, im_elem)
            insert_after(ref_para._element, sp_elem)
        
        print(f"  {fig_name}: {len(fig_paths)} figures after paragraph {para_idx}")
    
    doc.save(str(OUTPUT))
    print(f"\nSaved: {OUTPUT}")
    
    import zipfile
    zf = zipfile.ZipFile(str(OUTPUT))
    media = [n for n in zf.namelist() if 'media' in n]
    doc_xml = zf.read('word/document.xml').decode()
    om = doc_xml.count('m:oMath')
    print(f"  OMML equations: {om}")
    print(f"  Embedded images: {len(media)}")

if __name__ == "__main__":
    main()
