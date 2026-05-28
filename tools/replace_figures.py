"""
Replace all figures in paper_manuscript.docx with v4 versions.
Delete old image paragraphs, insert new ones at correct locations.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

FIG_DIR = Path("minimal_pinn") / "results" / "paper_figures" / "v4"
MORPH_DIR = Path("minimal_pinn") / "results" / "analysis" / "divergence_morphology_v1"

doc = Document("paper_manuscript.docx")

NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

# ── Step 1: Find and delete ALL existing image paragraphs ──
body = doc.element.body
removed = 0
for p_elem in list(body):
    drawings = p_elem.findall(".//" + f"{{{NSMAP['w']}}}drawing")
    if drawings:
        body.remove(p_elem)
        removed += 1
print(f"Removed {removed} image paragraphs")

# ── Step 2: Also clear orphan caption/figure paragraphs (短文本带"图") ──
for p_elem in list(body):
    texts = p_elem.findall(f".//{{{NSMAP['w']}}}t")
    full_text = "".join(t.text or "" for t in texts)
    if len(full_text) < 20 and "图" in full_text:
        body.remove(p_elem)
        removed += 1
print(f"Removed {removed} total paragraphs")

# ── Step 3: Find target paragraphs by text content ──
def find_para(text_fragment):
    """Find paragraph index containing fragment."""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return None

def insert_figure_after(doc, para_idx, fig_path, caption=""):
    """Insert figure (caption + image) after paragraph."""
    ref_elem = doc.paragraphs[para_idx]._element
    body = ref_elem.getparent()
    ref_pos = list(body).index(ref_elem)
    
    if not fig_path.exists():
        return False
    
    # Create caption paragraph
    cap_p = ref_elem.makeelement(f"{{{NSMAP['w']}}}p", {})
    cap_pPr = cap_p.makeelement(f"{{{NSMAP['w']}}}pPr", {})
    cap_jc = cap_pPr.makeelement(f"{{{NSMAP['w']}}}jc", {})
    cap_jc.set(f"{{{NSMAP['w']}}}val", "center")
    cap_pPr.append(cap_jc)
    cap_p.append(cap_pPr)
    cap_r = cap_p.makeelement(f"{{{NSMAP['w']}}}r", {})
    cap_rPr = cap_r.makeelement(f"{{{NSMAP['w']}}}rPr", {})
    cap_sz = cap_rPr.makeelement(f"{{{NSMAP['w']}}}sz", {})
    cap_sz.set(f"{{{NSMAP['w']}}}val", "18")  # 9pt
    cap_rPr.append(cap_sz)
    cap_r.append(cap_rPr)
    cap_t = cap_r.makeelement(f"{{{NSMAP['w']}}}t", {})
    cap_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    cap_t.text = caption
    cap_r.append(cap_t)
    cap_p.append(cap_r)
    
    # Create image paragraph using python-docx for proper image handling
    # Temporarily add to doc, get its XML, then move
    from docx import Document as Doc
    tmp = Doc()
    img_p_tmp = tmp.add_paragraph()
    img_p_tmp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = img_p_tmp.add_run()
    r.add_picture(str(fig_path), width=Inches(6.0))
    
    # Get the image paragraph's XML from the temp doc
    img_elem = img_p_tmp._element
    img_elem.getparent().remove(img_elem)  # detach from temp doc
    
    # Insert after reference
    body.insert(ref_pos + 1, cap_p)
    body.insert(ref_pos + 2, img_elem)
    return True

# ── Step 4: Insert all figures ──
inserts = [
    # Section 5.1: Case phase maps (one per case)
    ("**泊松方程。**", [("图5-1a 泊松方程：相对L2误差分布", FIG_DIR/"fig5-1a.png")]),
    ("**斯托克斯-泊肃叶流。**", [("图5-1b 斯托克斯-泊肃叶流：相对L2误差分布", FIG_DIR/"fig5-1b.png")]),
    ("**费希尔方程。**", [("图5-1c 费希尔方程：相对L2误差分布", FIG_DIR/"fig5-1c.png")]),
    ("**伯格斯方程。**", [("图5-1d 伯格斯方程：相对L2误差分布", FIG_DIR/"fig5-1d.png")]),
    # Section 5.2: Three-system probability heatmaps
    ("5.2 三系统的定量梯度", [
        ("图5-3a 斯托克斯-泊肃叶流：概率边界越界率", FIG_DIR/"fig5-3a.png"),
        ("图5-3b 费希尔方程：概率边界越界率", FIG_DIR/"fig5-3b.png"),
        ("图5-3c 伯格斯方程：概率边界越界率", FIG_DIR/"fig5-3c.png"),
        ("图5-4 边界关键点越界率的威尔逊置信区间", FIG_DIR/"fig5-4.png"),
    ]),
    # Section 6.2: Ablation
    ("6.2 消融分析", [
        ("图6-1 消融对比：完整四维R与简化指标的跨种子排序一致性", FIG_DIR/"fig6-1.png"),
    ]),
    # Section 6.3: Divergence morphology
    ("图6-2a", [
        ("图6-2a 伯格斯方程：仅由综合分识别为最差的工况（真值/预测/差值）", MORPH_DIR/"burgers_R-worst_not_L2-worst.png"),
    ]),
    ("图6-2b", [
        ("图6-2b 伯格斯方程：仅由相对误差识别为最差的工况（真值/预测/差值）", MORPH_DIR/"burgers_L2-worst_not_R-worst.png"),
    ]),
    # Section 7: Validation
    ("7.1 校准敏感性与反循环检查", [
        ("图7-1 校准敏感性：不同分位点配置下的主导维度分布", FIG_DIR/"fig7-1.png"),
        ("图7-2 反循环校验：留出评估中主导维度的一致率", FIG_DIR/"fig7-2.png"),
    ]),
    ("7.2 阈值可移植性与跨变体一致性", [
        ("图7-3 伯格斯方程：最安全点越界率随阈值倍数的变化", FIG_DIR/"fig7-3.png"),
    ]),
]

total = 0
for anchor_text, fig_list in inserts:
    pi = find_para(anchor_text)
    if pi is None:
        print(f"  SKIP: anchor not found: {anchor_text[:40]}")
        continue
    for caption, fig_path in fig_list:
        if insert_figure_after(doc, pi, fig_path, caption):
            total += 1
            print(f"  OK: {caption[:60]} after P{pi}")

doc.save("paper_manuscript.docx")
print(f"\nInserted {total} figures. Saved.")
