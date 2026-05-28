"""
Post-process pandoc-generated .docx: insert paper figures at the right locations.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PROJECT = Path(__file__).resolve().parent.parent
FIG_DIR_V3 = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v3"  
FIG_DIR_V2 = PROJECT / "minimal_pinn" / "results" / "paper_figures" / "v2"
MORPH_DIR = PROJECT / "minimal_pinn" / "results" / "analysis" / "divergence_morphology_v1"
INPUT = PROJECT / "paper_manuscript_pandoc.docx"
OUTPUT = PROJECT / "paper_manuscript_zh.docx"

FIGURES_POOL = {
    "图 1-(a)": FIG_DIR_V3 / "fig01_rel_l2_ab.png",
    "图 1-(b)": FIG_DIR_V3 / "fig01_rel_l2_cd.png",
    "图 2-(a)": FIG_DIR_V3 / "fig02_R_ab.png",
    "图 2-(b)": FIG_DIR_V3 / "fig02_R_cd.png",
    "图 3-(a)": FIG_DIR_V3 / "fig03_boundary_a.png",
    "图 3-(b)": FIG_DIR_V3 / "fig03_boundary_b.png",
    "图 3-(c)": FIG_DIR_V3 / "fig03_boundary_c.png",
    "图 4-(a)": FIG_DIR_V3 / "fig04_ablation_a.png",
    "图 4-(b)": FIG_DIR_V3 / "fig04_ablation_b.png",
    "图 4-(c)": FIG_DIR_V3 / "fig04_ablation_c.png",
    "图 5": MORPH_DIR / "burgers_R-worst_not_L2-worst.png",
    "图 6": MORPH_DIR / "burgers_L2-worst_not_R-worst.png",
    "图 A1": FIG_DIR_V2 / "fig_a1_calibration_sensitivity.png",
    "图 A2": FIG_DIR_V2 / "fig_a2_anti_circularity.png",
    "图 A3": FIG_DIR_V2 / "fig_a3_baseline_failure.png",
}

def insert_figure_after(paragraph, fig_path, caption=""):
    """Insert a figure (caption + image) after the given paragraph in the doc body."""
    body = paragraph._element.getparent()
    idx = list(body).index(paragraph._element)
    
    # Caption paragraph
    cp = paragraph._element.makeelement(qn("w:p"), {})
    cpPr = cp.makeelement(qn("w:pPr"), {})
    jc = cpPr.makeelement(qn("w:jc"), {})
    jc.set(qn("w:val"), "center")
    cpPr.append(jc)
    cp.append(cpPr)
    cr = cp.makeelement(qn("w:r"), {})
    crPr = cr.makeelement(qn("w:rPr"), {})
    sz = crPr.makeelement(qn("w:sz"), {})
    sz.set(qn("w:val"), "18")  # 9pt
    crPr.append(sz)
    cr.append(crPr)
    ct = cr.makeelement(qn("w:t"), {})
    ct.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    ct.text = caption
    cr.append(ct)
    cp.append(cr)
    
    # Image paragraph
    ip = paragraph._element.makeelement(qn("w:p"), {})
    ipPr = ip.makeelement(qn("w:pPr"), {})
    jc2 = ipPr.makeelement(qn("w:jc"), {})
    jc2.set(qn("w:val"), "center")
    ipPr.append(jc2)
    ip.append(ipPr)
    ir = ip.makeelement(qn("w:r"), {})
    idraw = ir.makeelement(qn("w:drawing"), {})
    
    # Build inline drawing
    with open(str(fig_path), "rb") as f:
        import base64
        img_data = base64.b64encode(f.read()).decode()
    
    from PIL import Image as PILImage
    img = PILImage.open(str(fig_path))
    w, h = img.size
    # Scale to 6 inches wide
    target_w_inch = 6.0
    target_h_inch = target_w_inch * h / w
    w_emu = int(target_w_inch * 914400)
    h_emu = int(target_h_inch * 914400)
    
    # Relationship ID for image
    rId = f"rId_fig_{hash(str(fig_path)) % 100000}"
    
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    
    inline = idraw.makeelement(f"{{{wp_ns}}}inline", {})
    
    extent = inline.makeelement(f"{{{wp_ns}}}extent", {})
    extent.set("cx", str(w_emu))
    extent.set("cy", str(h_emu))
    inline.append(extent)
    
    effectExtent = inline.makeelement(f"{{{wp_ns}}}effectExtent", {})
    effectExtent.set("l", "0"); effectExtent.set("t", "0")
    effectExtent.set("r", "0"); effectExtent.set("b", "0")
    inline.append(effectExtent)
    
    docPr = inline.makeelement(f"{{{wp_ns}}}docPr", {})
    docPr.set("id", str(hash(str(fig_path)) % 1000000))
    docPr.set("name", fig_path.name)
    inline.append(docPr)
    
    cNvGraphicFramePr = inline.makeelement(f"{{{wp_ns}}}cNvGraphicFramePr", {})
    inline.append(cNvGraphicFramePr)
    
    graphic = inline.makeelement(f"{{{a_ns}}}graphic", {})
    graphicData = graphic.makeelement(f"{{{a_ns}}}graphicData", {})
    graphicData.set("uri", pic_ns)
    
    pic = graphicData.makeelement(f"{{{pic_ns}}}pic", {})
    nvPicPr = pic.makeelement(f"{{{pic_ns}}}nvPicPr", {})
    cNvPr = nvPicPr.makeelement(f"{{{pic_ns}}}cNvPr", {})
    cNvPr.set("id", "0")
    cNvPr.set("name", fig_path.name)
    nvPicPr.append(cNvPr)
    cNvPicPr = nvPicPr.makeelement(f"{{{pic_ns}}}cNvPicPr", {})
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)
    
    blipFill = pic.makeelement(f"{{{pic_ns}}}blipFill", {})
    blip = blipFill.makeelement(f"{{{a_ns}}}blip", {})
    blip.set(f"{{{r_ns}}}embed", rId)
    blipFill.append(blip)
    stretch = blipFill.makeelement(f"{{{a_ns}}}stretch", {})
    fillRect = stretch.makeelement(f"{{{a_ns}}}fillRect", {})
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)
    
    spPr = pic.makeelement(f"{{{pic_ns}}}spPr", {})
    xfrm = spPr.makeelement(f"{{{a_ns}}}xfrm", {})
    off = xfrm.makeelement(f"{{{a_ns}}}off", {})
    off.set("x", "0"); off.set("y", "0")
    xfrm.append(off)
    ext2 = xfrm.makeelement(f"{{{a_ns}}}ext", {})
    ext2.set("cx", str(w_emu)); ext2.set("cy", str(h_emu))
    xfrm.append(ext2)
    spPr.append(xfrm)
    prstGeom = spPr.makeelement(f"{{{a_ns}}}prstGeom", {})
    prstGeom.set("prst", "rect")
    spPr.append(prstGeom)
    pic.append(spPr)
    
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    idraw.append(inline)
    ir.append(idraw)
    ip.append(ir)
    
    # Spacer paragraph
    sp = paragraph._element.makeelement(qn("w:p"), {})
    
    # Insert after the target paragraph
    insert_pos = idx + 1
    body.insert(insert_pos, cp)
    body.insert(insert_pos + 1, ip) 
    body.insert(insert_pos + 2, sp)
    
    return rId

def main():
    doc = Document(str(INPUT))
    
    # Track figures to insert
    to_insert = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        for fig_name, fig_path in FIGURES_POOL.items():
            if fig_name in text and fig_path.exists() and fig_name not in [x[1] for x in to_insert]:
                to_insert.append((i, fig_name, fig_path))
    
    # Sort by paragraph index (descending so we can insert without shifting)
    to_insert.sort(key=lambda x: -x[0])
    
    # Build image relationships
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    image_part = None
    rId_counter = 1
    
    for para_idx, fig_name, fig_path in to_insert:
        if para_idx >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[para_idx]
        
        # Add image to docx (just use add_picture approach directly)
        from docx.opc.part import Part
        # Simpler: just add a paragraph with the picture
        from docx.shared import Inches as DocxInches
        """
        
        # Calculate size
        from PIL import Image as PILImage
        img = PILImage.open(str(fig_path))
        w, h = img.size
        target_w = Inches(6.0)
        target_h = Inches(6.0 * h / w)
        
        # Add a new paragraph after the current one with the image
        # Since python-docx doesn't have insert_paragraph_after, 
        # we use the underlying XML
        body = p._element.getparent()
        idx = list(body).index(p._element)
        
        # Caption
        from docx.shared import Pt as PtShared
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(fig_name)
        cap_run.font.size = PtShared(9)
        cap_run.font.name = "宋体"
        cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        cap_elem = cap_para._element
        body.remove(cap_elem)
        body.insert(idx + 1, cap_elem)
        
        # Image
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_para.add_run()
        img_run.add_picture(str(fig_path), width=target_w)
        img_elem = img_para._element
        body.remove(img_elem)
        body.insert(idx + 2, img_elem)
        
        # Spacer
        sp_para = doc.add_paragraph()
        sp_elem = sp_para._element
        body.remove(sp_elem)
        body.insert(idx + 3, sp_elem)
        
        print(f"  Inserted {fig_name} after paragraph {para_idx}")
    
    doc.save(str(OUTPUT))
    print(f"\nSaved: {OUTPUT}")
    print(f"  Figures inserted: {len(to_insert)}")

if __name__ == "__main__":
    main()
